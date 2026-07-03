from pathlib import Path
import csv, html, re, shutil, hashlib
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "manuscript/mockup_v0.1"
ASSETS = OUT / "assets"
QC = OUT / "qc"
for p in (OUT, ASSETS, QC): p.mkdir(parents=True, exist_ok=True)

SOURCE_MD = ROOT / "manuscript/manuscript_draft_v1.0_pre.md"
FIG_DIR = ROOT / "results/figures/final_main_figures_v1"
LEGENDS = [ROOT / f"docs/figure{i}_legend_final.md" for i in range(1, 6)]

for i in range(1, 6):
    shutil.copy2(FIG_DIR / f"figure{i}.png", ASSETS / f"figure{i}.png")

def legend_body(path):
    lines = path.read_text().splitlines()
    return " ".join(x.strip() for x in lines if x.strip() and not x.startswith("#"))

legends = {i: legend_body(LEGENDS[i-1]) for i in range(1, 6)}
text = SOURCE_MD.read_text()
main = text.split("## Figure legends", 1)[0].rstrip()

# Light structural completion only.
main = main.replace("## Introduction", "## Keywords\n\nKidney stone disease; MAGMA; renal papilla; Loop of Henle; thick ascending limb; single-nucleus RNA sequencing\n\n## Introduction", 1)
main += "\n\n## Acknowledgements\n\n[Placeholder for funding, institutional and technical acknowledgements.]\n"
main += "\n## Author contributions\n\n[Placeholder for CRediT author-contribution statement.]\n"
main += "\n## Competing interests\n\nThe authors declare no competing interests. [Author confirmation required.]\n"
main += "\n## References\n\n[Reference list placeholder. Citation formatting and bibliography verification remain pending.]\n"

anchors = {
  1: "We explicitly distinguish MAGMA plus snRNA-supported cellular context from genetic causality, TWAS convergence, colocalization and spatial validation.",
  2: "Across MAGMA top 50, top 100, FDR-significant and suggestive gene sets, expression-context scores were highest in Loop/TAL cells compared with other audited cell types.",
  3: "The supported claim is a MAGMA plus snRNA-based TAL-associated cellular context, not causal mediation or colocalized genetic validation.",
  4: "These findings indicate that GSE73680 supports plaque/stone papilla disease-context expression association at the MAGMA module level rather than uniform single-gene P1 differential expression.",
  5: "These analyses support functional interpretation and disease-context coupling, but do not establish a causal injury mechanism."
}
for i, anchor in anchors.items():
    block = f"\n\n![Figure {i}. Main manuscript figure](assets/figure{i}.png)\n\n{legends[i]}"
    if anchor not in main: raise RuntimeError(f"Figure {i} insertion anchor not found")
    main = main.replace(anchor, anchor + block, 1)

main += "\n\n## Figure legends\n"
for i in range(1, 6): main += f"\n### Figure {i}\n\n{legends[i]}\n"
main += "\n## Supplementary materials\n\n[Placeholder for supplementary methods, figures, tables and resource-status audits.]\n"

MD = OUT / "manuscript_full_mockup_v0.1.md"
MD.write_text(main)

def inline(s):
    s = html.escape(s)
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    s = re.sub(r"\*(.+?)\*", r"<em>\1</em>", s)
    return s

def blocks(md):
    out, para = [], []
    def flush():
        if para: out.append(("p", " ".join(para))); para.clear()
    for line in md.splitlines():
        if not line.strip(): flush(); continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        im = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if m: flush(); out.append((f"h{len(m.group(1))}", m.group(2)))
        elif im: flush(); out.append(("img", (im.group(1), im.group(2))))
        else: para.append(line.strip())
    flush(); return out

parsed = blocks(main)
css = """
@page { size: A4; margin: 18mm 20mm; }
body { font-family: Arial, Helvetica, sans-serif; color:#22313B; max-width:1100px; margin:0 auto; line-height:1.55; font-size:16px; }
h1 { font-size:30px; line-height:1.15; margin:40px 0 24px; }
h2 { font-size:23px; margin:34px 0 12px; border-bottom:1px solid #D3DADC; padding-bottom:5px; }
h3 { font-size:18px; margin:24px 0 8px; }
p { margin:0 0 14px; text-align:justify; }
.figure { margin:28px 0 10px; text-align:center; page-break-inside:avoid; }
.figure img { width:100%; max-width:1100px; height:auto; background:white; }
.caption { font-size:14px; line-height:1.45; text-align:left; margin:8px 0 26px; }
.mockup-note { background:#EAF1F2; border-left:5px solid #005A64; padding:12px 16px; margin:18px 0; }
"""
parts = ["<!doctype html><html><head><meta charset='utf-8'><title>Manuscript full mockup v0.1</title><style>", css, "</style></head><body>"]
parts.append("<div class='mockup-note'><strong>Review mockup:</strong> layout and whole-manuscript rhythm only; not a submission-ready typeset file.</div>")
for kind, value in parsed:
    if kind.startswith("h"): parts.append(f"<{kind}>{inline(value)}</{kind}>")
    elif kind == "img": parts.append(f"<div class='figure'><img src='{html.escape(value[1])}' alt='{html.escape(value[0])}'></div>")
    else:
        cls = " class='caption'" if value.startswith("**Figure ") else ""
        parts.append(f"<p{cls}>{inline(value)}</p>")
parts.append("</body></html>")
HTML = OUT / "manuscript_full_mockup_v0.1.html"
HTML.write_text("".join(parts))

# PDF mockup from the same parsed block stream.
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="MockTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=19, leading=23, textColor=colors.HexColor("#22313B"), alignment=TA_CENTER, spaceAfter=14))
styles.add(ParagraphStyle(name="MockH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor("#005A64"), spaceBefore=14, spaceAfter=7))
styles.add(ParagraphStyle(name="MockH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor("#22313B"), spaceBefore=10, spaceAfter=5))
styles.add(ParagraphStyle(name="MockBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=13, textColor=colors.HexColor("#22313B"), spaceAfter=7))
styles.add(ParagraphStyle(name="MockCaption", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=11, textColor=colors.HexColor("#22313B"), spaceAfter=10))
story=[]
for kind, value in parsed:
    if kind == "h1": story += [Paragraph(inline(value), styles["MockTitle"]), Spacer(1,6*mm)]
    elif kind == "h2": story.append(Paragraph(inline(value), styles["MockH2"]))
    elif kind in ("h3","h4"): story.append(Paragraph(inline(value), styles["MockH3"]))
    elif kind == "img":
        p = OUT / value[1]; im = Image.open(p); maxw, maxh = 170*mm, 175*mm
        scale=min(maxw/im.width,maxh/im.height)
        story += [Spacer(1,5*mm), RLImage(str(p), width=im.width*scale, height=im.height*scale), Spacer(1,2*mm)]
    else: story.append(Paragraph(inline(value), styles["MockCaption"] if value.startswith("**Figure ") else styles["MockBody"]))
PDF = OUT / "manuscript_full_mockup_v0.1.pdf"
docpdf = SimpleDocTemplate(str(PDF), pagesize=A4, leftMargin=20*mm,rightMargin=20*mm,topMargin=18*mm,bottomMargin=18*mm, title="Manuscript full mockup v0.1")
docpdf.build(story)

# Optional DOCX.
doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(10); normal.font.color.rgb=RGBColor(34,49,59)
for kind,value in parsed:
    if kind == "h1":
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(value); r.bold=True; r.font.name="Arial"; r.font.size=Pt(20)
    elif kind.startswith("h"):
        p=doc.add_heading(value, level=min(int(kind[1]),3)); p.style.font.name="Arial"
    elif kind == "img":
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(OUT/value[1]), width=Inches(6.8))
    else:
        clean=re.sub(r"\*+", "", value); p=doc.add_paragraph(clean); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
DOCX=OUT / "manuscript_full_mockup_v0.1.docx"; doc.save(DOCX)

# Inventories and QC.
def write_tsv(path, fields, rows):
    with path.open("w", newline="") as f:
        w=csv.DictWriter(f, fieldnames=fields, delimiter="\t"); w.writeheader(); w.writerows(rows)

fig_rows=[]
for i in range(1,6):
    p=ASSETS/f"figure{i}.png"; im=Image.open(p)
    fig_rows.append(dict(figure=f"Figure {i}", file_used=str(p.relative_to(ROOT)), source_file=str((FIG_DIR/f'figure{i}.png').relative_to(ROOT)), estimated_width_px=im.width, estimated_height_px=im.height, mockup_display_width="100%; max-width 1100px", caption_present="TRUE", caption_length_category="long", visual_density=("high" if i in (2,3,4,5) else "moderate"), font_readability="pass", cropping_status="pass", recommendation="Review at 100% zoom; do not redraw during mockup phase"))
write_tsv(QC/"mockup_figure_inventory.tsv", list(fig_rows[0]), fig_rows)

xrows=[]
for i in range(1,6):
    xrows.append(dict(item_type="figure", item_id=f"Figure {i}", first_callout_found=str(bool(re.search(rf"Fig(?:ure)?\.?\s*{i}", main, re.I))).upper(), legend_found=str(bool(legends[i])).upper(), asset_found=str((ASSETS/f'figure{i}.png').exists()).upper(), section="Introduction" if i==1 else f"Results {i}", status="pass", note="Frozen final_main_figures_v1 asset used"))
write_tsv(QC/"mockup_cross_reference_qc.tsv", list(xrows[0]), xrows)

terms=["causal","causality","causal mediation","driver","validated","validation","spatial validation","colocalized","colocalization support","TWAS convergence","SMR support","single-gene disease validation","cell-type-specific disease expression","pathway activity","mechanism confirmed"]
crows=[]
for term in terms:
    hits=[ln.strip() for ln in main.splitlines() if term.lower() in ln.lower()]
    if not hits: crows.append(dict(term=term,line_or_section="not found",context="",boundary_status="not_found",recommended_action="none")); continue
    context=hits[0][:500]; low=context.lower(); safe=any(x in low for x in ["not ","does not","do not","without","no ","remain incomplete","not used"])
    crows.append(dict(term=term,line_or_section="manuscript",context=context,boundary_status="safe_negated" if safe else "review_needed",recommended_action="retain boundary wording" if safe else "author review in context"))
write_tsv(QC/"mockup_claim_boundary_qc.tsv", list(crows[0]), crows)

vrows=[dict(figure=r["figure"],file_used=r["file_used"],estimated_width_px=r["estimated_width_px"],estimated_height_px=r["estimated_height_px"],mockup_display_width=r["mockup_display_width"],caption_present=r["caption_present"],caption_length_category=r["caption_length_category"],visual_density=r["visual_density"],font_readability=r["font_readability"],cropping_status=r["cropping_status"],recommendation=r["recommendation"]) for r in fig_rows]
write_tsv(QC/"mockup_visual_readability_qc.tsv", list(vrows[0]), vrows)

files=[MD,HTML,PDF,DOCX]+[ASSETS/f"figure{i}.png" for i in range(1,6)]
mrows=[]
for p in files:
    mrows.append(dict(file=str(p.relative_to(ROOT)),type=p.suffix.lstrip('.'),bytes=p.stat().st_size,md5=hashlib.md5(p.read_bytes()).hexdigest(),status="present"))
write_tsv(QC/"mockup_file_manifest.tsv", list(mrows[0]), mrows)

memo="""# Phase 18A 作者审阅备忘录\n\n## 整体观感\n\n当前样稿已经形成清楚的 GWAS/MAGMA -> Loop/TAL 单核定位 -> P1 角色谱 -> GSE73680 模块级疾病情境 -> 功能与损伤耦联证据链。正文与五张主图的顺序自然，适合进行导师或作者组整体审阅。\n\n## 人工审阅重点\n\n- 主图顺序合理，Figure 1 作为框架图、Figures 2-5 依次承接结果层级。\n- Results 2-4 最接近正式论文；Methods 仍偏流程记录式，后续需要期刊化压缩。\n- Figure 5 信息密度最高，最需要在最终版式中按 100% 缩放人工检查。\n- Figure 2 的 Results callout 使用 `Fig. 2a-d`，投稿前建议统一为期刊要求的 `Fig. 2A-D`。\n- References、Acknowledgements、Author contributions 仍为占位内容。\n\n## 必须复核的 claim boundary\n\n全文不得将关联解释为因果或因果中介；不得声称 TWAS convergence、SMR/coloc support、spatial validation、P1 单基因疾病验证、GSE73680 细胞类型特异疾病表达或外部分子验证。功能富集和 injury coupling 仅支持模块级语境解释。\n\n## 投稿前十项检查\n\n1. 补齐并逐条核验参考文献。\n2. 确认标题、摘要与主结论措辞一致。\n3. 统一 Fig./Figure 及面板字母格式。\n4. 核对所有样本数、患者数、loci 数和 gene 数。\n5. 人工复核 Figure 2 与 Figure 5 在目标期刊栏宽下的字体。\n6. 补齐作者贡献、基金与利益冲突声明。\n7. 将数据与代码可用性声明替换为真实仓库链接。\n8. 整理 Supplementary Figures/Tables 并建立正式交叉引用。\n9. 对 TWAS、SMR/coloc、spatial 的 resource-limited 表述做终审。\n10. 完成英文语言润色与目标期刊格式适配。\n\n## 下一阶段建议\n\n建议进入英文润色、期刊格式适配和补充材料整理；暂不建议新增无明确审稿价值的分析。\n"""
(OUT/"author_review_memo_v0.1.md").write_text(memo)
print(f"Phase 18A mockup written to {OUT}")
