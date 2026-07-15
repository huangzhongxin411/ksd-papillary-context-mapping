#!/usr/bin/env python3
"""Create the targeted v3.3 manuscript text and layout-preserving DOCX."""

from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
V2_MD = ROOT / "manuscript/manuscript_v3.2_final_polished.md"
V3_MD = ROOT / "manuscript/manuscript_v3.3_language_polished.md"
V2_DOCX = ROOT / "manuscript/manuscript_v3.2_final_polished.docx"
V3_DOCX = ROOT / "manuscript/manuscript_v3.3_language_polished.docx"


REPLACEMENTS = {
    "Post-GWAS renal papillary context mapping of kidney stone disease genetic risk":
        "Post-GWAS mapping of kidney stone disease genetic risk to renal papillary contexts",

    "Background: Kidney stone disease (KSD) has a substantial polygenic component, but the renal cell and tissue contexts in which associated genes are represented remain incompletely resolved. We integrated trans-ancestry KSD genome-wide association summary statistics with renal papillary single-nucleus, spatial and paired bulk transcriptomic resources. MAGMA supplied an EUR-linkage-disequilibrium-reference-based genetic-priority layer, while GTEx Kidney_Cortex S-PrediXcan was treated as proxy annotation rather than papilla-specific regulatory evidence.":
        "Background: Kidney stone disease (KSD) has a substantial polygenic component, but the renal cell and tissue contexts of associated genes remain incompletely resolved. We integrated trans-ancestry KSD genome-wide association summary statistics with renal papillary single-nucleus, spatial and paired bulk transcriptomic resources. MAGMA defined an EUR-linkage-disequilibrium-reference-based genetic-priority layer, and GTEx Kidney_Cortex S-PrediXcan provided proxy annotation.",

    "Results: Quality control retained 4,915,033 of 5,960,489 GWAS rows, including 3,209 genome-wide significant variants, and identified 60 lead variants in 57 reconstructed loci. MAGMA tested 17,316 genes and identified 94 Bonferroni-significant, 369 false-discovery-rate (FDR)-significant and 187 suggestive genes. In GSE231569, 43,878 nuclei from four papillary donors, including 540 Loop of Henle/thick ascending limb (Loop/TAL) nuclei, supported a donor-level Loop/TAL-associated context in expression- and detection-matched random benchmarks, with attenuation but retention of the broader pattern after known-driver removal. Spatial label transfer across ten complete sections, comprising four GSE206306 sections and six GSE231630 sections, was usable only for broad-compartment projection; Loop/TAL predictions were zero or sparse, precluding claim-grade Loop/TAL or lesion enrichment. Kidney_Cortex S-PrediXcan tested 5,989 genes and yielded 51 FDR-supported models, of which 42 were one-SNP and nine were multi-SNP models. A repaired 235-gene reporting set comprised R1-R6 groups of 68, 1, 2, 141, 16 and 7 genes, respectively. In GSE73680, 55 samples from 29 patients, including 26 paired patients, showed concordant positive module coefficients that did not remain FDR-significant and were attenuated in sensitivity analyses using a KIM1/LCN2-like injury proxy and related tissue-state scores.":
        "Results: Quality control retained 4,915,033 of 5,960,489 GWAS rows, including 3,209 genome-wide significant variants, and identified 60 lead variants in 57 reconstructed loci. MAGMA tested 17,316 genes and identified 94 Bonferroni-significant, 369 false-discovery-rate (FDR)-significant and 187 suggestive genes. In GSE231569, 43,878 nuclei from four papillary donors, including 540 Loop of Henle/thick ascending limb (Loop/TAL) nuclei, supported a donor-level Loop/TAL-associated context in expression- and detection-matched random benchmarks; the broader pattern attenuated but remained after known-driver removal. Across ten complete spatial sections, comprising four GSE206306 sections and six GSE231630 sections, label transfer supported broad-compartment projection but not Loop/TAL or lesion enrichment. Kidney_Cortex S-PrediXcan tested 5,989 genes and yielded 51 FDR-supported models, including 42 one-SNP and nine multi-SNP models. The evidence-stratified 235-gene reporting set comprised R1-R6 groups of 68, 1, 2, 141, 16 and 7 genes, respectively. In GSE73680, 55 samples from 29 patients, including 26 paired patients, showed concordant positive module coefficients that did not remain FDR-significant and attenuated in sensitivity analyses using a KIM1/LCN2-like injury proxy and related tissue-state scores.",

    "Conclusions: Genetically prioritized KSD genes map to a donor-level Loop/TAL-associated renal papillary context and are embedded in a broader KIM1/LCN2-like injury-associated bulk disease context. These results provide post-GWAS functional interpretation and context mapping, not causal gene validation, causal cell-type assignment, plaque-specific localization or papilla-specific regulatory inference.":
        "Conclusions: Genetically prioritized KSD genes map to a donor-level Loop/TAL-associated renal papillary context and are embedded in a broader KIM1/LCN2-like injury-associated bulk disease context. This post-GWAS framework provides cellular and tissue-context interpretation while leaving causal genes, causal cell types, plaque localization and papilla-specific regulation unresolved.",

    "Here, we developed an evidence-stratified post-GWAS framework to determine whether MAGMA-prioritized KSD genes map to coherent renal papillary contexts. We separated EUR-LD-reference-based MAGMA genetic priority from Kidney_Cortex TWAS proxy annotation, evaluated donor-level single-nucleus patterns in GSE231569, tested paired bulk module responses in GSE73680 with tissue-state sensitivity analyses, and used four GSE206306 and six GSE231630 sections only for descriptive broad-compartment spatial projection. This design links genetic priority to cellular and tissue context while preserving the distinct strength and limitations of each evidence layer.":
        "Here, we developed an evidence-stratified post-GWAS framework to ask whether MAGMA-prioritized KSD genes map to coherent renal papillary contexts. We distinguished EUR-LD-reference-based genetic priority from Kidney_Cortex TWAS proxy annotation, evaluated donor-level single-nucleus patterns, examined paired bulk module responses with tissue-state sensitivity analyses, and used spatial data for supplementary broad-compartment projection. This design connects genetic priority to cellular and tissue contexts while preserving the distinct inferential role of each evidence layer.",

    "Auditable GWAS and MAGMA analyses define an EUR-LD-reference-based genetic-priority layer":
        "GWAS and MAGMA analyses define an EUR-LD-reference-based genetic-priority layer",
    "The locked GWAS input contained 5,960,489 rows. Quality control retained 4,915,033 rows, including 3,209 variants at genome-wide significance. Distance pruning identified 60 lead SNPs, and merging overlapping windows produced 57 reconstructed loci (Supplementary Table 1). The source publication reported 59 loci,[5] but the publication-level lead-SNP or locus table required to reconcile the two counts was not available. We therefore report the 57-locus result as a reproducible reconstruction rather than as a replacement for the published locus definition.":
        "The GWAS input contained 5,960,489 rows. Quality control retained 4,915,033 rows, including 3,209 variants at genome-wide significance. Figure 1 provides a quality-control Manhattan visualization of the cleaned summary statistics, with genome-wide significant variants retained; locus-level and MAGMA-ranked gene information is provided in Supplementary Figure S1 and Supplementary Table 1. Distance pruning identified 60 lead SNPs, and merging overlapping windows produced 57 reconstructed loci. The source publication reported 59 loci,[5] but the publication-level lead-SNP or locus table required to reconcile the two counts was unavailable. We therefore report the 57-locus result as a reproducible reconstruction rather than a replacement for the published locus definition.",
    "MAGMA v1.10 tested 17,316 genes using GRCh37 gene locations and the 1000 Genomes European reference panel. Ninety-four genes met Bonferroni significance, 369 met FDR significance and 187 met the prespecified suggestive threshold of P < 1 x 10^-4 (Supplementary Table 2). Locked top-50, top-100, Bonferroni, FDR and suggestive modules were propagated without post hoc membership changes. Because the source GWAS was trans-ancestry whereas the linkage-disequilibrium reference was European, these results define an EUR-LD-reference-based genetic-priority layer, not ancestry-generalizable fine mapping or causal-gene evidence (Supplementary Figure S1; Supplementary Table 2).":
        "MAGMA v1.10 tested 17,316 genes using GRCh37 gene locations and the 1000 Genomes European reference panel. Ninety-four genes met Bonferroni significance, 369 met FDR significance and 187 met the prespecified suggestive threshold of P < 1 x 10^-4 (Supplementary Figure S1; Supplementary Table 2). The top-50, top-100, Bonferroni, FDR and suggestive modules were defined before downstream scoring and retained without post hoc membership changes. Because the source GWAS was trans-ancestry whereas the linkage-disequilibrium reference was European, these results constitute an EUR-LD-reference-based genetic-priority layer rather than ancestry-generalizable fine mapping or causal-gene evidence.",

    "The audited GSE231569 papillary single-nucleus dataset contained 43,878 nuclei from four donors. Harmonized broad-compartment annotation identified 540 nuclei in the Loop/TAL compartment. Donor-by-compartment summaries placed the locked MAGMA modules consistently toward the Loop/TAL end of the expression distribution, with donors serving as the biological support units rather than individual nuclei (Figure 2; Supplementary Table 3).":
        "The GSE231569 papillary single-nucleus dataset contained 43,878 nuclei from four donors. Harmonized broad-compartment annotation identified 540 nuclei in the Loop/TAL compartment. Donor-by-compartment summaries placed the prespecified MAGMA modules consistently toward the Loop/TAL end of the expression distribution, with donors serving as the biological support units rather than individual nuclei (Figure 2; Supplementary Table 3).",

    "Locked MAGMA module scores were overlaid descriptively on the tissue sections and interpreted only against broad transferred compartments (Supplementary Figure S2). These overlays show where module signal occurs within the assayed papillary tissue, but they do not localize genetic risk to plaque or infer disease-control differences, papilla-specific regulation or causal anatomy.":
        "MAGMA module scores were overlaid descriptively on the tissue sections and interpreted only against broad transferred compartments (Supplementary Figure S2). These overlays provide supplementary tissue-context projection within the assayed papillary sections; the absence of lesion regions and usable Loop/TAL transfer precludes plaque localization, disease-control comparison or papilla-specific regulatory inference.",

    "A repaired candidate model separates MAGMA priority from Kidney_Cortex proxy annotation":
        "An evidence-stratified candidate model separates MAGMA priority from Kidney_Cortex proxy annotation",
    "S-PrediXcan with GTEx v8 Kidney_Cortex models tested 5,989 genes and identified 51 FDR-supported models. Forty-two were one-SNP models and nine were multi-SNP models, making the evidence predominantly weak and variant-limited. Kidney_Cortex also does not match the renal papilla. TWAS was therefore retained as proxy annotation and was not used to upgrade genes to papilla-specific regulatory or causal status (Supplementary Figure S3; Supplementary Table 4).":
        "S-PrediXcan with GTEx v8 Kidney_Cortex models tested 5,989 genes and identified 51 FDR-supported models. Forty-two were one-SNP models and nine were multi-SNP models, making support predominantly variant-limited. Because Kidney_Cortex does not match the renal papilla, TWAS was retained as proxy annotation and did not confer papilla-specific regulatory or causal status (Supplementary Figure S3; Supplementary Table 4).",
    "The repaired candidate universe contained 235 genes assigned deterministically to six mutually exclusive reporting groups: R1, 68; R2, 1; R3, 2; R4, 141; R5, 16; and R6, 7 (Figure 3; Supplementary Table 5). These are reporting groups, not causal tiers. R1-R6 organize combinations of MAGMA priority, proxy TWAS support and contextual review; curated exemplars remain interpretive annotations and do not alter assignment. Bulk status indicates paired disease-context and tissue-state review and is not used to upgrade candidate reporting groups. Complete resources for claim-grade SMR/HEIDI or colocalization were unavailable, and no such result was used.":
        "The candidate universe contained 235 genes assigned deterministically to six mutually exclusive reporting groups: R1, 68; R2, 1; R3, 2; R4, 141; R5, 16; and R6, 7 (Figure 3; Supplementary Figure S3; Supplementary Table 5). R1-R6 are reporting groups, not causal tiers. They organize combinations of MAGMA priority, proxy TWAS support and contextual review; curated exemplars remain interpretive annotations and do not alter assignment. Bulk status records paired disease-context and tissue-state review without upgrading a reporting group. Complete resources for formal SMR/HEIDI or colocalization were unavailable, and no such result was used.",

    "The reconstructed GSE73680 dataset comprised 55 papillary samples from 29 patients, including 26 patients with paired control/adjacent and plaque/stone-associated tissue. Group labels were filename-derived and manually reviewable. Module scores were calculated as the arithmetic mean of gene-wise z scores across detected members, without an additional log transformation. In patient fixed-effect models, all five locked MAGMA modules had positive group coefficients (0.169-0.174), but none remained significant after FDR correction (FDR 0.071-0.082) (Figure 4; Supplementary Table 6).":
        "The reconstructed GSE73680 dataset comprised 55 papillary samples from 29 patients, including 26 patients with paired control/adjacent and plaque/stone-associated tissue. Group labels were derived from filenames, a metadata limitation considered in interpretation. Module scores were calculated as the arithmetic mean of gene-wise z scores across detected members, without an additional log transformation. In patient fixed-effect models, all five prespecified MAGMA modules had positive group coefficients (0.169-0.174), but none remained significant after FDR correction (FDR 0.071-0.082) (Figure 4; Supplementary Table 6).",
    "GSE73680 feature-level expression and filename-derived, manually reviewable group labels yielded 55 samples from 29 patients, including 26 complete pairs.[12,37] Module scores were arithmetic means of gene-wise z scores across the 55 samples for detected module genes. Paired deltas were defined as plaque/stone-associated minus control/adjacent, and patient fixed-effect models used `module_score ~ group_binary + factor(patient_id)`. Benjamini-Hochberg correction was applied across the five canonical modules.":
        "GSE73680 feature-level expression and filename-derived group labels yielded 55 samples from 29 patients, including 26 complete pairs.[12,37] Filename-based metadata provenance was retained as an interpretive limitation. Module scores were arithmetic means of gene-wise z scores across the 55 samples for detected module genes. Paired deltas were defined as plaque/stone-associated minus control/adjacent, and patient fixed-effect models used `module_score ~ group_binary + factor(patient_id)`. Benjamini-Hochberg correction was applied across the five canonical modules.",

    "The principal contribution of this study is an evidence-stratified map connecting KSD gene-based association signals to renal papillary transcriptomic contexts. Rather than treating overlap across datasets as validation, the analysis assigns each layer a defined inferential role. MAGMA anchors genetic priority, donor-level snRNA identifies a Loop/TAL-associated context, spatial data provide descriptive anatomical orientation, Kidney_Cortex TWAS supplies proxy annotation and paired bulk data capture the disease-associated tissue background.":
        "This study provides an evidence-stratified map connecting KSD gene-based association signals to renal papillary transcriptomic contexts. The strongest contextual evidence arose from donor-level snRNA data, which placed prioritized modules in a Loop/TAL-associated setting. MAGMA anchored genetic priority, spatial data supplied supplementary anatomical orientation, Kidney_Cortex TWAS contributed proxy annotation and paired bulk data placed the modules within a disease-associated tissue background. Assigning these distinct roles prevents cross-dataset agreement from being interpreted as equivalent evidence.",
    "The Loop/TAL-associated single-nucleus pattern is biologically compatible with renal transport and ion-handling pathways implicated in stone susceptibility. Established genes including *UMOD*, *CASR*, *CLDN14* and *CLDN10* provide useful biological anchors.[19-24] Their familiarity, however, cannot identify the gene mediating a GWAS locus. Expression- and detection-matched random benchmarks supported the donor-level Loop/TAL-minus-other contrast across canonical modules. Known-driver removal attenuated but did not abolish the broader module-level pattern; this sensitivity does not establish independence from Loop/TAL biology. The four-donor design and one donor with only four Loop/TAL nuclei appropriately limit the result to moderate contextual support.":
        "The Loop/TAL-associated pattern is biologically compatible with renal transport and ion-handling pathways implicated in stone susceptibility. Established genes including *UMOD*, *CASR*, *CLDN14* and *CLDN10* provide biological anchors,[19-24] although their familiarity cannot identify the gene mediating a GWAS locus. Expression- and detection-matched random benchmarks supported the donor-level Loop/TAL-minus-other contrast across canonical modules, and known-driver removal attenuated but did not abolish the broader pattern. Together, these findings provide moderate contextual support; the four-donor design, including one donor with only four Loop/TAL nuclei, limits stronger generalization and does not establish independence from Loop/TAL biology.",
    "The bulk results provide a complementary tissue-state view. Concordant positive coefficients across five modules were not FDR-significant and were attenuated by KIM1/LCN2-like injury-proxy and epithelial-state sensitivity analyses. This attenuation is informative because it places genetically prioritized programs within a broader papillary response rather than isolating a disease-specific module. It does not distinguish transcriptional regulation within cells from shifts in cellular composition.":
        "The bulk results provide a complementary tissue-state view. Concordant positive coefficients across five modules, although not FDR-significant, placed genetically prioritized programs within a broader papillary response. Their attenuation in KIM1/LCN2-like injury-proxy and epithelial-state sensitivity analyses supports tissue-state embedding rather than an isolated disease-specific module. Because these expression scores are not cell fractions, the data cannot distinguish within-cell transcriptional regulation from changes in tissue composition.",
    "The spatial and TWAS layers chiefly define boundaries. Ten spatial sections from GSE206306 and GSE231630 permitted broad-compartment projection, but sparse or absent Loop/TAL transfer and missing lesion annotations prevented plaque- or Loop/TAL-specific inference. Likewise, the predominance of one-SNP Kidney_Cortex models and the cortex-to-papilla mismatch prevented regulatory upgrading of individual genes. Neither missing claim-grade SMR/colocalization nor weak proxy support should be interpreted as evidence against biological relevance; they identify experiments and resources still needed.":
        "Spatial and TWAS evidence added complementary but more limited context. Ten spatial sections from GSE206306 and GSE231630 supported broad-compartment projection, whereas sparse or absent Loop/TAL transfer and missing lesion annotations prevented plaque- or Loop/TAL-specific inference. The predominance of one-SNP Kidney_Cortex models and the cortex-to-papilla mismatch similarly restricted TWAS to proxy annotation. The absence of complete SMR/HEIDI or colocalization inputs does not argue against biological relevance; it identifies the locus-resolved and papilla-relevant resources needed for stronger regulatory interpretation.",

    "The use of a 1000 Genomes European LD reference for a trans-ancestry GWAS limits ancestry-generalizable interpretation. The 57 reconstructed loci could not be reconciled with the 59 published loci without the source lead-SNP or locus table. The single-nucleus analysis contained four donors and only 540 Loop/TAL nuclei, and one donor contributed only four Loop/TAL nuclei. Matched-random benchmarks supported the canonical contrasts, whereas known-driver removal attenuated but did not abolish the broader pattern. Spatial projection was descriptive, included two lower-confidence sections and lacked usable Loop/TAL prediction and lesion annotations. Kidney_Cortex TWAS was tissue-mismatched and dominated by one-SNP models, while complete inputs for claim-grade SMR/HEIDI and colocalization were unavailable. The bulk dataset was reconstructed from microarray data, and its group labels were filename-derived and manually reviewable; its tissue-state scores were not cell fractions, and module associations did not survive FDR correction. All analyses were observational, and no perturbational experiment tested a prioritized gene, module or cell context.":
        "The genetic analyses have two principal limitations. A 1000 Genomes European LD reference was used for a trans-ancestry GWAS, limiting ancestry-generalizable interpretation, and the 57 reconstructed loci could not be reconciled with the 59 published loci without the source lead-SNP or locus table.\n\nThe cell- and tissue-context analyses were also constrained by sampling and resolution. The single-nucleus dataset contained four donors and 540 Loop/TAL nuclei, with one donor contributing only four Loop/TAL nuclei; matched-random support attenuated after known-driver removal. Spatial projection was descriptive, included two lower-confidence sections and lacked usable Loop/TAL prediction and lesion annotations.\n\nKidney_Cortex TWAS was tissue-mismatched and dominated by one-SNP models, while complete inputs for formal SMR/HEIDI and colocalization were unavailable. The bulk dataset was reconstructed from microarray data with filename-derived group labels; its tissue-state scores were not cell fractions, and module associations did not survive FDR correction. All analyses were observational, and perturbational experiments will be required to test prioritized genes, modules and cellular contexts.",

    "MAGMA v1.10 used NCBI build 37 gene locations and the 1000 Genomes European LD reference.[25] Of 17,316 tested genes, significance was defined using Bonferroni correction, Benjamini-Hochberg FDR,[32] and a suggestive threshold of P < 1 x 10^-4. Top-50, top-100, Bonferroni, FDR and suggestive lists were frozen before downstream scoring. Executable, version, gene-location, LD-prefix, SNP/P input, gene output and log provenance were audited. MAGMA results were interpreted as EUR-LD-reference-based genetic priority.":
        "MAGMA v1.10 used NCBI build 37 gene locations and the 1000 Genomes European LD reference.[25] Of 17,316 tested genes, significance was defined using Bonferroni correction, Benjamini-Hochberg FDR,[32] and a suggestive threshold of P < 1 x 10^-4. Top-50, top-100, Bonferroni, FDR and suggestive lists were defined before downstream scoring. Executable, version, gene-location, LD-prefix, SNP/P input, gene output and log provenance were recorded for reproducibility. MAGMA results were interpreted as EUR-LD-reference-based genetic priority.",
    "Four papillary donors from GSE231569 were retained, yielding 43,878 audited nuclei after quality control and broad-compartment harmonization. Seurat was used for single-nucleus object handling, normalization and broad-compartment workflows.[33] The Loop/TAL compartment contained 540 nuclei. For each frozen module, nucleus-level scores were calculated as arithmetic means of log-normalized expression among detected genes and summarized by donor and compartment. Donors were biological support units. Leave-one-donor-out summaries, 1,000 expression- and detection-matched random gene sets per module, low-cell-count sensitivity, and panel-level removal of prespecified Loop/TAL/transport drivers evaluated donor dependence, matched expectation and driver dependence.":
        "Four papillary donors from GSE231569 were retained, yielding 43,878 nuclei after quality control and broad-compartment harmonization. Seurat was used for single-nucleus object handling, normalization and broad-compartment workflows.[33] The Loop/TAL compartment contained 540 nuclei. For each prespecified module, nucleus-level scores were calculated as arithmetic means of log-normalized expression among detected genes and summarized by donor and compartment. Donors were biological support units. Leave-one-donor-out summaries, 1,000 expression- and detection-matched random gene sets per module, low-cell-count sensitivity, and panel-level removal of prespecified Loop/TAL/transport drivers evaluated donor dependence, matched expectation and driver dependence.",
    "Ten complete sections with expression, coordinate and image assets were processed as Seurat spatial objects, comprising four GSE206306 and six GSE231630 sections.[33] The GSE231569 reference supported label transfer only at the broad-compartment level. Transfer confidence and prediction sparsity were audited by section. Because transferred Loop/TAL prediction scores were zero or nonzero-sparse and no lesion-region annotations were available, module scores were displayed as descriptive overlays without formal Loop/TAL, plaque, disease-control or cross-section enrichment tests. Sections, not spots, defined the highest available biological unit.":
        "Ten complete sections with expression, coordinate and image assets were processed as Seurat spatial objects, comprising four GSE206306 and six GSE231630 sections.[33] The GSE231569 reference supported label transfer only at the broad-compartment level. Transfer confidence and prediction sparsity were assessed by section. Because transferred Loop/TAL prediction scores were zero or nonzero-sparse and no lesion-region annotations were available, module scores were displayed as descriptive overlays without formal Loop/TAL, plaque, disease-control or cross-section enrichment tests. Sections, not spots, defined the highest available biological unit.",
    "S-PrediXcan used GTEx v8 Kidney_Cortex MASHR models.[29-30] Benjamini-Hochberg correction was applied across 5,989 tested genes, and 51 FDR-supported models were classified by the number of contributing variants as one-SNP (42) or multi-SNP (9). The tissue was treated as a cortex proxy, and TWAS did not upgrade causal or papilla-specific regulatory status. A deterministic union and deduplication procedure produced 235 candidate genes and mutually exclusive R1-R6 reporting groups (68, 1, 2, 141, 16 and 7 genes). The groups are reporting categories, not causal tiers. Feasibility of SMR/HEIDI[34] and colocalization[35-36] was audited, but incomplete inputs precluded claim-grade analyses.":
        "S-PrediXcan used GTEx v8 Kidney_Cortex MASHR models.[29-30] Benjamini-Hochberg correction was applied across 5,989 tested genes, and 51 FDR-supported models were classified by the number of contributing variants as one-SNP (42) or multi-SNP (9). The tissue was treated as a cortex proxy, and TWAS did not confer causal or papilla-specific regulatory status. A deterministic union and deduplication procedure produced 235 candidate genes and mutually exclusive R1-R6 reporting groups (68, 1, 2, 141, 16 and 7 genes). The groups are reporting categories, not causal tiers. Feasibility of SMR/HEIDI[34] and colocalization[35-36] was evaluated, but incomplete inputs precluded formal analyses.",
    "Tests were two-sided where applicable, with correction families defined within each analysis layer. Donors, patients and sections, rather than nuclei, samples or spots, were treated as biological support units for cross-unit interpretation. Machine-readable source data, manifests, logs and scripts were retained for reported tables and figures. No exploratory SMR, colocalization, plaque localization or spatial Loop/TAL analysis was promoted to claim-grade evidence.":
        "Tests were two-sided where applicable, with correction families defined within each analysis layer. Donors, patients and sections, rather than nuclei, samples or spots, were treated as biological support units for cross-unit interpretation. Machine-readable source data, manifests, logs and scripts were retained for reported tables and figures. No SMR, colocalization, plaque-localization or spatial Loop/TAL result was included among the reported evidence.",

    "The Cao et al. KSD GWAS summary statistics are available from Zenodo (https://doi.org/10.5281/zenodo.14790324). The associated GWAS Catalog study record is GCST90652506. Public transcriptomic data are available from the Gene Expression Omnibus under accessions GSE231569, GSE73680, GSE206306 and GSE231630. GTEx v8 / PredictDB Kidney_Cortex MASHR models are available from their original distribution resource. Analysis code, derived result tables, figure Source Data, Supplementary Tables, manifests and selected logs are available at https://github.com/huangzhongxin411/ksd-papillary-context-mapping and will be archived at Zenodo upon release (Zenodo DOI to be added upon release).":
        "The Cao et al. KSD GWAS summary statistics are available from Zenodo (https://doi.org/10.5281/zenodo.14790324), and the associated GWAS Catalog study record is GCST90652506. Public transcriptomic data are available from the Gene Expression Omnibus under accessions GSE231569, GSE73680, GSE206306 and GSE231630. GTEx v8/PredictDB Kidney_Cortex MASHR models are available from their original distribution resource. Analysis code, derived result tables, figure Source Data, Supplementary Tables, manifests and selected logs are publicly available at https://github.com/huangzhongxin411/ksd-papillary-context-mapping, including GitHub Release v1.0.0. Zenodo archival of the manuscript-synchronized repository and its DOI remain pending and will be added after minting.",

    "Figure 1. GWAS quality-control Manhattan plot":
        "Figure 1. GWAS quality-control Manhattan plot",
    "Downsampled Manhattan representation of the locked trans-ancestry KSD GWAS after quality control, with all sampled genome-wide significant variants retained for visualization. The plot documents GWAS input behavior only. MAGMA gene counts and rank summaries are reported in the Results, Supplementary Figure S1 and Supplementary Tables; they are not displayed in this panel. The figure does not imply ancestry-generalizable fine mapping or causal genes.":
        "GWAS quality-control Manhattan plot generated from downsampled cleaned trans-ancestry KSD summary statistics, with genome-wide significant variants retained. The horizontal reference line marks genome-wide significance. The panel visualizes GWAS input behavior and is not a fine-mapping or causal-gene inference display. Locus-level and MAGMA-ranked gene information is provided in Supplementary Figure S1 and Supplementary Table 1. File: Figure_1_GWAS_QC_Manhattan_polished.",
    "Figure 2. Donor-level snRNA evidence for a Loop/TAL-associated context":
        "Figure 2. Donor-level snRNA context of genetically prioritized modules",
    "Donor-aware module summaries in 43,878 nuclei from four GSE231569 papillary donors, including 540 Loop/TAL nuclei. Panels show broad-compartment context, donor-level module contrasts, matched-random benchmarking, leave-one-donor-out sensitivity and known-driver removal. Donors are the biological support units. Expression- and detection-matched random benchmarks support the donor-level contrast, whereas known-driver removal attenuates but does not abolish the broader pattern. The figure supports a moderate Loop/TAL-associated context, not causal cell-type assignment or independence from Loop/TAL biology.":
        "Donor-aware analysis of 43,878 nuclei from four GSE231569 papillary donors, including 540 Loop/TAL nuclei. (A) snRNA resource summary. (B) Broad-compartment UMAP. (C) Loop/TAL marker-expression context. (D) Donor-by-compartment scores for prespecified MAGMA modules. (E) Expression- and detection-matched random benchmark for the Loop/TAL-minus-other contrast. (F) Sensitivity after removal of prespecified Loop/TAL and renal-transport drivers. Donor x compartment is the biological unit. The benchmark and driver-removal analyses support a moderate Loop/TAL-associated context and do not establish a causal cell type. File: Figure_2_snRNA_context_polished.",
    "Figure 3. Repaired candidate reporting model":
        "Figure 3. Candidate reporting model",
    "Evidence organization for 235 genes across six mutually exclusive reporting groups: R1 = 68, R2 = 1, R3 = 2, R4 = 141, R5 = 16 and R6 = 7. MAGMA priority is separated from GTEx Kidney_Cortex proxy TWAS support, including one-SNP and multi-SNP model classes. R1-R6 are reporting groups rather than causal tiers, and curated exemplars do not alter group assignment. Bulk status indicates paired disease-context and tissue-state review and is not used to upgrade candidate reporting groups.":
        "Evidence-stratified organization of the candidate reporting model. (A) Parallel roles of genetic priority, cell context, proxy annotation, tissue context and disease-state review. (B) Representative rows from the candidate evidence matrix. (C) Kidney_Cortex proxy-model burden, including one-SNP and multi-SNP classes. (D) Counts across R1-R6. (E) Interpretation boundary. R1-R6 are reporting groups, not causal tiers. Kidney_Cortex TWAS is proxy annotation, and contextual evidence does not upgrade causal status. The complete 235-gene matrix is shown in Supplementary Figure S3 and Supplementary Table 5. File: Figure_3_candidate_reporting_model_polished.",
    "Figure 4. Paired bulk disease-context analysis":
        "Figure 4. Paired bulk disease-context and tissue-state analysis",
    "GSE73680 analysis of 55 samples from 29 patients, including 26 complete pairs. Five canonical MAGMA module coefficients were positive but did not remain FDR-significant. Tissue-state and paired-delta sensitivity analyses show attenuation and association with a broader KIM1/LCN2-like injury-proxy background. Group labels were filename-derived and manually reviewable; marker scores are expression proxies, not cell fractions.":
        "Paired GSE73680 analysis of 55 samples from 29 patients, including 26 complete pairs. (A) Compact paired-study design. (B) Selected paired module scores. (C) Base disease-context coefficients for prespecified MAGMA modules. (D) Paired tissue-state proxy shifts. (E) Coupling between module and tissue-state paired differences. (F) Attenuation of disease-context coefficients after tissue-state adjustment. Positive module coefficients did not remain FDR-significant and attenuated in sensitivity analyses, placing prioritized modules in a broader KIM1/LCN2-like injury-associated tissue state. Expression scores are not cell fractions, and the analysis does not provide causal validation. File: Figure_4_bulk_disease_context_polished.",

    "Supplementary Figure S1. GWAS and MAGMA diagnostic plots":
        "Supplementary Figure S1. GWAS and MAGMA diagnostics",
    "Quantile-quantile diagnostics, P-value distribution and MAGMA gene-rank summaries complement the main-text GWAS QC Manhattan plot. These panels document input behavior and gene-priority thresholds; they do not provide locus fine mapping or causal-gene inference.":
        "GWAS and MAGMA diagnostic panels complement the main-text quality-control Manhattan plot. The panels show GWAS quantile-quantile and P-value behavior together with MAGMA significance and rank summaries; the top 20 MAGMA-ranked genes are displayed in panel S1C. These genes represent gene-based priority and should not be interpreted as causal genes. File: Supplementary_Figure_S1_GWAS_MAGMA_diagnostics_polished.",
    "Supplementary Figure S2. Descriptive spatial broad-compartment projection":
        "Supplementary Figure S2. Spatial broad-compartment tissue-context projection",
    "Broad-compartment label-transfer quality control and descriptive locked-module overlays across ten sections comprising four GSE206306 and six GSE231630 sections. Two sections showed lower confidence, Loop/TAL predictions were zero or sparse and no lesion-region annotation was available. The panels do not support plaque localization, Loop/TAL enrichment, disease-control comparison or causal spatial inference.":
        "Five-page supplementary spatial projection across ten complete sections, comprising four GSE206306 and six GSE231630 sections. S2A presents spatial-resource and label-transfer quality control. S2B1-S2B3 show Bonferroni, Top 100 and suggestive module overlays, and S2B4 summarizes section-level module scores by predicted broad compartment. The transferred labels provide broad tissue context only: no lesion ROI was available, Loop/TAL prediction was unusable, and no Loop/TAL enrichment, disease-control test or plaque localization was performed. File: Supplementary_Figure_S2_spatial_projection_polished.",
    "Supplementary Figure S3. Full Kidney_Cortex TWAS proxy and candidate evidence display":
        "Supplementary Figure S3. Kidney_Cortex TWAS proxy and candidate evidence matrix",
    "Complete S-PrediXcan and candidate evidence summaries for 5,989 tested models, including 51 FDR-supported results (42 one-SNP and nine multi-SNP models), together with the 235-gene reporting universe. The figure presents cortex-proxy annotation and reporting categories, not papilla-specific regulation or causal tiers.":
        "Thirteen-page supplement separating TWAS proxy quality from the complete candidate matrix. S3A summarizes 5,989 tested Kidney_Cortex models and 51 FDR-supported results, including 42 one-SNP and nine multi-SNP models. S3B paginates the complete 235-gene matrix by R group, with every cell rendered directly from Supplementary Table 5. R1-R6 are reporting groups, not causal tiers, and Kidney_Cortex TWAS remains proxy annotation rather than papilla-specific regulatory evidence. File: Supplementary_Figure_S3_TWAS_candidate_proxy_polished.",
}


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS.items():
        source = old
        target = new
        for prefix in ("Background:", "Results:", "Conclusions:"):
            if old.startswith(prefix):
                source = f"**{prefix}**" + old[len(prefix):]
                target = f"**{prefix}**" + new[len(prefix):]
                break
        if source not in text:
            raise RuntimeError(f"Expected source text not found: {old[:100]}")
        text = text.replace(source, target, 1)
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if "\n\n" in text:
        raise ValueError("Multi-paragraph replacement requires XML insertion")
    first = paragraph.runs[0] if paragraph.runs else None
    font_snapshot = None
    if first is not None:
        font_snapshot = {
            "name": first.font.name, "size": first.font.size, "bold": first.bold,
            "italic": first.italic, "underline": first.underline,
        }
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    prefix = next((p for p in ("Background:", "Results:", "Conclusions:") if text.startswith(p)), None)
    if prefix:
        run = paragraph.add_run(prefix)
        run.bold = True
        paragraph.add_run(text[len(prefix):])
    else:
        run = paragraph.add_run(text)
        if font_snapshot:
            run.font.name = font_snapshot["name"]
            run.font.size = font_snapshot["size"]
            run.bold = font_snapshot["bold"]
            run.italic = font_snapshot["italic"]
            run.underline = font_snapshot["underline"]


def insert_paragraph_after(paragraph, text: str):
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[paragraph._parent._element.body.index(new_p)] if False else None
    # python-docx does not expose an insertion constructor; wrap through the paragraph class.
    from docx.text.paragraph import Paragraph
    wrapped = Paragraph(new_p, paragraph._parent)
    set_paragraph_text(wrapped, text)
    return wrapped


def replace_docx_text(doc: Document) -> None:
    pending = dict(REPLACEMENTS)
    # The markdown title omits the leading '#'; all prose values match DOCX paragraphs.
    for paragraph in list(doc.paragraphs):
        old = paragraph.text
        key = old if old in pending else next((candidate for candidate in pending if candidate.replace("*", "").replace("`", "") == old), None)
        if key is None:
            continue
        new = pending.pop(key).replace("*", "").replace("`", "")
        if "\n\n" not in new:
            set_paragraph_text(paragraph, new)
            continue
        pieces = new.split("\n\n")
        set_paragraph_text(paragraph, pieces[0])
        anchor = paragraph
        for piece in pieces[1:]:
            anchor = insert_paragraph_after(anchor, piece)
    # Values that are headings identical in old/new may not be distinguishable as a change.
    pending = {old: new for old, new in pending.items() if old != new}
    if pending:
        raise RuntimeError("DOCX replacements not found: " + "; ".join(k[:80] for k in pending))


def replace_figure_images(doc: Document) -> None:
    images = [
        ROOT / "figures_v3.3/Figure_1_GWAS_QC_Manhattan_polished.png",
        ROOT / "figures_v3.3/Figure_2_snRNA_context_polished.png",
        ROOT / "figures_v3.3/Figure_3_candidate_reporting_model_polished.png",
        ROOT / "figures_v3.3/Figure_4_bulk_disease_context_polished.png",
    ]
    drawing_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
    if len(drawing_paragraphs) != 4:
        raise RuntimeError(f"Expected four embedded main figures, found {len(drawing_paragraphs)}")
    from PIL import Image
    max_width = 5_897_880
    for index, (paragraph, image_path) in enumerate(zip(drawing_paragraphs, images), 1):
        blip = paragraph._p.xpath(".//a:blip")[0]
        rid = blip.get(qn("r:embed"))
        part = doc.part.related_parts[rid]
        part._blob = image_path.read_bytes()
        with Image.open(image_path) as image:
            width, height = image.size
        cx = max_width
        cy = int(cx * height / width)
        for extent in paragraph._p.xpath(".//wp:extent") + paragraph._p.xpath(".//a:xfrm/a:ext"):
            extent.set("cx", str(cx))
            extent.set("cy", str(cy))
        for doc_pr in paragraph._p.xpath(".//wp:docPr"):
            doc_pr.set("name", f"Figure {index}")
            doc_pr.set("descr", image_path.stem)


def synchronize_document_furniture(doc: Document) -> None:
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                set_paragraph_text(paragraph, "BMC Genomics manuscript | Language-polished v3.3 draft")
        for paragraph in section.footer.paragraphs:
            for text_node in paragraph._p.xpath(".//w:t"):
                if text_node.text and "Manuscript v3.2 final polished" in text_node.text:
                    text_node.text = text_node.text.replace("Manuscript v3.2 final polished", "Manuscript v3.3 language-polished draft")
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.style.name.startswith("Heading") and re.match(r"^(?:Supplementary )?Figure ", paragraph.text):
            paragraphs[index + 1].alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_change_log() -> None:
    path = ROOT / "manuscript/manuscript_v3.3_revision_change_log.md"
    path.write_text("""# Manuscript v3.3 revision change log

## Scope
This version applies targeted language polish and synchronizes the manuscript with the Phase 7-Step 2 figure package. Manuscript v3.2 remains unchanged. No analysis, accepted value, threshold, group assignment or scientific conclusion was changed.

## Language changes
- Revised the title for a clearer post-GWAS context-mapping cadence without strengthening the claim.
- Reordered the abstract around the genetic-priority anchor, donor-level snRNA result, supporting spatial/TWAS layers and bulk tissue-state result.
- Replaced internal workflow vocabulary such as `locked`, `audited`, `repaired`, `reviewable` and `claim-grade` with journal-facing alternatives where scientific provenance did not require it.
- Smoothed transitions across the six Results subsections while preserving their architecture.
- Rebalanced the Discussion around finding, interpretation and evidence-layer limitation.
- Grouped Limitations into genetic, cell/spatial and TWAS/bulk constraints.
- Clarified that GitHub Release v1.0.0 is public while repository Zenodo archival and DOI remain pending.

## Figure synchronization
- Replaced the four embedded v3.2 main figures with the approved v3.3 polished PNGs while retaining the manuscript layout.
- Updated Figure 1 to remain exclusively a GWAS quality-control Manhattan visualization; no gene labels were added.
- Synchronized Figure 2 panels A-F, Figure 3 panels A-E and Figure 4 panels A-F with the polished files.
- Updated Supplementary Figure S2 as a five-page spatial supplement and Supplementary Figure S3 as a 13-page TWAS/candidate supplement.
- Retained Supplementary Table 5 as the authoritative complete 235-gene record.

## Preserved boundaries
- MAGMA defines EUR-LD-reference-based genetic priority, not causal genes.
- snRNA provides donor-level Loop/TAL-associated context, not causal cell-type assignment.
- Spatial data provide supplementary broad tissue context without lesion ROI or Loop/TAL enrichment.
- Kidney_Cortex TWAS remains proxy annotation.
- Bulk scores support disease-context and tissue-state embedding and are not cell fractions.
- R1-R6 remain reporting groups, not causal tiers.

## Human-only placeholders retained
- Institutional confirmation of ethics wording.
- CRediT author contributions.
- Final funding and competing-interest confirmation.
- All-author approval.
- Repository Zenodo DOI after minting.
""", encoding="utf-8")


def main() -> None:
    source_md = V2_MD.read_text(encoding="utf-8")
    V3_MD.write_text(apply_replacements(source_md), encoding="utf-8")
    shutil.copy2(V2_DOCX, V3_DOCX)
    doc = Document(V3_DOCX)
    replace_docx_text(doc)
    replace_figure_images(doc)
    synchronize_document_furniture(doc)
    doc.core_properties.title = "Post-GWAS mapping of kidney stone disease genetic risk to renal papillary contexts"
    doc.core_properties.subject = "Research article manuscript v3.3 language-polished draft"
    doc.core_properties.comments = "Targeted language polish and figure-legend synchronization; scientific content locked to v3.2."
    doc.save(V3_DOCX)
    build_change_log()
    print("Created v3.3 Markdown, DOCX, and revision change log.")


if __name__ == "__main__":
    main()
