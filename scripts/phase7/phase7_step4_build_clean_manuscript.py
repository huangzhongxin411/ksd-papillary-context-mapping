#!/usr/bin/env python3
"""Create the v3.3 clean submission manuscript without internal draft furniture."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "manuscript/manuscript_v3.3_language_polished.md"
SOURCE_DOCX = ROOT / "manuscript/manuscript_v3.3_language_polished.docx"
TARGET_MD = ROOT / "manuscript/manuscript_v3.3_submission_clean.md"
TARGET_DOCX = ROOT / "manuscript/manuscript_v3.3_submission_clean.docx"

OLD_AVAILABILITY = (
    "The Cao et al. KSD GWAS summary statistics are available from Zenodo "
    "(https://doi.org/10.5281/zenodo.14790324), and the associated GWAS Catalog study record is "
    "GCST90652506. Public transcriptomic data are available from the Gene Expression Omnibus under "
    "accessions GSE231569, GSE73680, GSE206306 and GSE231630. GTEx v8/PredictDB Kidney_Cortex MASHR "
    "models are available from their original distribution resource. Analysis code, derived result tables, "
    "figure Source Data, Supplementary Tables, manifests and selected logs are publicly available at "
    "https://github.com/huangzhongxin411/ksd-papillary-context-mapping, including GitHub Release v1.0.0. "
    "Zenodo archival of the manuscript-synchronized repository and its DOI remain pending and will be "
    "added after minting."
)

NEW_AVAILABILITY = (
    "The Cao et al. KSD GWAS summary statistics are available from Zenodo "
    "(https://doi.org/10.5281/zenodo.14790324), and the associated GWAS Catalog study record is "
    "GCST90652506. Public transcriptomic data are available from the Gene Expression Omnibus under "
    "accessions GSE231569, GSE73680, GSE206306 and GSE231630. GTEx v8/PredictDB Kidney_Cortex MASHR "
    "models are available from their original distribution resource. Analysis code, derived result tables, "
    "figure source data, Supplementary Tables, manifests and selected logs are available at the GitHub "
    "repository (https://github.com/huangzhongxin411/ksd-papillary-context-mapping). The current "
    "manuscript-synchronized release will be v1.0.1 after final approval; GitHub Release v1.0.0 is retained "
    "as a pre-Zenodo checkpoint. A manuscript-synchronized repository archive and DOI will be added after "
    "final release and Zenodo minting."
)


def strip_file_note(text: str) -> str:
    return re.sub(r"\s+File:\s+(?:Figure|Supplementary_Figure)_[A-Za-z0-9_.-]+\.$", "", text)


def set_paragraph_text(paragraph, text: str) -> None:
    first = paragraph.runs[0] if paragraph.runs else None
    font = None
    if first is not None:
        font = (first.font.name, first.font.size, first.bold, first.italic, first.underline)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if font:
        run.font.name, run.font.size, run.bold, run.italic, run.underline = font


def build_markdown() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    if OLD_AVAILABILITY not in text:
        raise RuntimeError("Expected v3.3 Data Availability paragraph not found in Markdown")
    text = text.replace(OLD_AVAILABILITY, NEW_AVAILABILITY, 1)
    text = re.sub(r"\s+File:\s+(?:Figure|Supplementary_Figure)_[A-Za-z0-9_.-]+\.(?=\n|$)", "", text)
    TARGET_MD.write_text(text, encoding="utf-8")


def build_docx() -> None:
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(TARGET_DOCX)
    availability_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text == OLD_AVAILABILITY:
            set_paragraph_text(paragraph, NEW_AVAILABILITY)
            availability_found = True
        elif " File: " in paragraph.text and re.match(r"^(?:GWAS|Donor-aware|Evidence-stratified|Paired GSE73680|GWAS and MAGMA|Five-page|Thirteen-page)", paragraph.text):
            set_paragraph_text(paragraph, strip_file_note(paragraph.text))
    if not availability_found:
        raise RuntimeError("Expected v3.3 Data Availability paragraph not found in DOCX")

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            set_paragraph_text(paragraph, "")
        for paragraph in section.footer.paragraphs:
            for text_node in paragraph._p.xpath(".//w:t"):
                if text_node.text and "draft" in text_node.text.lower():
                    text_node.text = "Page "

    doc.core_properties.subject = "Research article submission manuscript v3.3"
    doc.core_properties.comments = "Clean submission manuscript; repository DOI pending human-approved release and Zenodo minting."
    doc.save(TARGET_DOCX)


def main() -> None:
    build_markdown()
    build_docx()
    print("Created clean v3.3 Markdown and DOCX.")


if __name__ == "__main__":
    main()
