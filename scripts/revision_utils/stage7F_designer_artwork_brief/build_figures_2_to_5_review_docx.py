#!/usr/bin/env python3
"""Build a visually verified Word review pack for Figures 2-5 and legends."""

from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / "docs/revision/stage7F_designer_artwork_brief/figures_2_to_5_review_with_legends_v1.0.docx"

FIGURES = [
    {
        "number": 2,
        "version": "draft v0.2",
        "title": "Donor-level snRNA context mapping of MAGMA-prioritized modules",
        "image": ROOT / "results/figures/revision/stage4C2R_draft_figures_v0.2/figure2_snRNA_context_draft_v0.2.png",
        "legend_file": ROOT / "docs/revision/stage4C2R_draft_figures_v0.2/draft_figure2_figure3_legends_v0.2.md",
    },
    {
        "number": 3,
        "version": "draft v0.2",
        "title": "Two-axis candidate-gene evidence model and curated exemplar boundaries",
        "image": ROOT / "results/figures/revision/stage4C2R_draft_figures_v0.2/figure3_candidate_evidence_draft_v0.2.png",
        "legend_file": ROOT / "docs/revision/stage4C2R_draft_figures_v0.2/draft_figure2_figure3_legends_v0.2.md",
    },
    {
        "number": 4,
        "version": "draft v0.1",
        "title": "Attenuated injury/remodeling-associated paired bulk context",
        "image": ROOT / "results/figures/revision/stage5C1_gse73680_figure4_draft/figure4_gse73680_bulk_context_draft_v0.1.png",
        "legend_file": ROOT / "docs/revision/stage5C1_gse73680_figure4_draft/draft_figure4_legend_v0.1.md",
    },
    {
        "number": 5,
        "version": "draft v0.1",
        "title": "Supplementary spatial and TWAS renal papillary context",
        "image": ROOT / "results/figures/revision/stage6C_spatial_twas_figure5_draft/figure5_spatial_twas_context_draft_v0.1.png",
        "legend_file": ROOT / "docs/revision/stage6C_spatial_twas_figure5_draft/draft_figure5_legend_v0.1.md",
    },
]

NAVY = RGBColor(36, 90, 100)
MUTED = RGBColor(98, 107, 110)
DARK = RGBColor(45, 52, 54)
GOLD = RGBColor(185, 155, 90)


def set_run_font(run, size: float, bold: bool = False, italic: bool = False,
                 color: RGBColor = DARK, name: str = "Arial") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, NAVY),
        ("Heading 2", 13, 14, 7, NAVY),
        ("Heading 3", 12, 10, 5, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Legend" not in styles:
        legend = styles.add_style("Figure Legend", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legend = styles["Figure Legend"]
    legend.font.name = "Arial"
    legend._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    legend._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    legend.font.size = Pt(10.5)
    legend.font.color.rgb = DARK
    legend.paragraph_format.space_before = Pt(0)
    legend.paragraph_format.space_after = Pt(6)
    legend.paragraph_format.line_spacing = 1.17

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("Figures 2-5 | Review copy"), 8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("Page "), 8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Figures 2-5 Review Pack"), 26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Current draft artwork with corresponding editable legends"), 13, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Prepared for visual and wording review | 30 June 2026"), 10, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.6)
    p.paragraph_format.right_indent = Inches(1.6)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run(
            "This document reorganizes the existing Figure 2-5 drafts and their current legends only. "
            "No figure values, panel content, evidence classification or legend wording has been scientifically revised."
        ),
        11,
        color=DARK,
    )

    for item in FIGURES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(2.2)
        run = p.add_run(f"Figure {item['number']}  ")
        set_run_font(run, 11, bold=True, color=NAVY)
        set_run_font(p.add_run(f"{item['version']} | {item['title']}"), 11, color=DARK)

    doc.add_page_break()


def extract_legend(item: dict) -> tuple[str, str]:
    text = item["legend_file"].read_text(encoding="utf-8")
    number = item["number"]
    if number in (2, 3):
        match = re.search(
            rf"^## Figure {number}\.\s+(.+?)\n\n(.+?)(?=\n## Figure|\Z)",
            text,
            flags=re.MULTILINE | re.DOTALL,
        )
        if not match:
            raise ValueError(f"Could not extract Figure {number} legend")
        return match.group(1).strip(), match.group(2).strip()

    paragraph = next((line.strip() for line in text.splitlines() if line.strip().startswith("**Figure")), "")
    if not paragraph:
        raise ValueError(f"Could not extract Figure {number} legend")
    title_match = re.match(r"\*\*Figure\s+\d+\s*\|\s*(.+?)\.\*\*\s*(.*)", paragraph)
    if not title_match:
        raise ValueError(f"Could not parse Figure {number} legend title")
    return title_match.group(1).strip(), title_match.group(2).strip()


def add_markdown_runs(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]), 10.5)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), 10.5, bold=True)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), 10.2, name="Courier New")
        else:
            set_run_font(paragraph.add_run(token[1:-1]), 10.5, italic=True)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), 10.5)


def add_figure_page(doc: Document, item: dict) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(3)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.add_run(f"Figure {item['number']}. {item['title']}")

    source = doc.add_paragraph()
    source.paragraph_format.space_after = Pt(4)
    set_run_font(source.add_run(f"Current artwork: {item['version']} | review copy; source-data locked"), 8.5, color=MUTED)

    with Image.open(item["image"]) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    max_width = 9.55
    max_height = 6.25
    width_in = min(max_width, max_height * aspect)
    height_in = width_in / aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(item["image"]), width=Inches(width_in), height=Inches(height_in))
    shape._inline.docPr.set("descr", f"Current draft Figure {item['number']}: {item['title']}")

    doc.add_page_break()


def add_legend_page(doc: Document, item: dict) -> None:
    legend_title, legend_body = extract_legend(item)
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(10)
    heading.add_run(f"Figure {item['number']} legend")

    p = doc.add_paragraph(style="Figure Legend")
    lead = p.add_run(f"Figure {item['number']} | {legend_title}. ")
    set_run_font(lead, 10.5, bold=True, color=NAVY)
    add_markdown_runs(p, legend_body)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("Review note: "), 9, bold=True, color=GOLD)
    set_run_font(
        note.add_run("Edit wording directly in this legend page; preserve sample units, numerical anchors and claim boundaries."),
        9,
        color=MUTED,
    )


def main() -> None:
    for item in FIGURES:
        if not item["image"].exists() or not item["legend_file"].exists():
            raise FileNotFoundError(item)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    for index, item in enumerate(FIGURES):
        add_figure_page(doc, item)
        add_legend_page(doc, item)
        if index < len(FIGURES) - 1:
            doc.add_page_break()

    doc.core_properties.title = "Figures 2-5 Review Pack with Legends"
    doc.core_properties.subject = "Current manuscript figure review"
    doc.core_properties.author = "Research manuscript team"
    doc.core_properties.keywords = "Figures 2-5, legends, review copy"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
