#!/usr/bin/env python3
from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
INPUT_MD = ROOT / "manuscript/manuscript_v3.1_targeted_revision.md"
BASE_DOCX = ROOT / "manuscript/manuscript_v3.1_targeted_revision.docx"
OUT_MD = ROOT / "manuscript/manuscript_v3.2_final_polished.md"
OUT_DOCX = ROOT / "manuscript/manuscript_v3.2_final_polished.docx"
REF_AUDIT = ROOT / "results/tables/phase6_step6_reference_renumbering_audit.tsv"
FIG1_AUDIT = ROOT / "results/tables/phase6_step6_figure1_callout_fix_audit.tsv"


def heading_at(text, pos):
    section = "Front matter"
    line_number = text[:pos].count("\n") + 1
    for match in re.finditer(r"(?m)^(#{2,3}) (.+)$", text[:pos]):
        section = match.group(2)
    return f"{section}, source line {line_number}"


def expand_citation(content):
    values = []
    for token in content.replace(" ", "").split(","):
        if not token:
            continue
        if "-" in token:
            start, end = map(int, token.split("-"))
            values.extend(range(start, end + 1))
        else:
            values.append(int(token))
    return values


def compress_citation(values):
    out = []
    start = previous = values[0]
    for value in values[1:]:
        if value == previous + 1:
            previous = value
            continue
        out.append(str(start) if start == previous else f"{start}-{previous}")
        start = previous = value
    out.append(str(start) if start == previous else f"{start}-{previous}")
    return ",".join(out)


def split_reference_entry(line):
    match = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
    if not match:
        raise ValueError(f"Malformed reference entry: {line[:80]}")
    return int(match.group(1)), match.group(2)


def short_title(entry):
    parts = entry.split(". ")
    title = parts[1] if len(parts) > 1 else parts[0]
    return title[:110]


def renumber_references(md):
    before_refs, remainder = md.split("## References", 1)
    refs_text, after_refs = remainder.split("## Figures and legends", 1)
    entries = {}
    for line in refs_text.strip().splitlines():
        if not line.strip():
            continue
        number, entry = split_reference_entry(line)
        entries[number] = entry

    first_order = []
    first_location = {}
    citation_pattern = re.compile(r"\[([0-9,\- ]+)\]")
    for match in citation_pattern.finditer(before_refs):
        for old_number in expand_citation(match.group(1)):
            if old_number not in first_order:
                first_order.append(old_number)
                first_location[old_number] = heading_at(before_refs, match.start())

    missing = sorted(set(first_order) - set(entries))
    if missing:
        raise ValueError(f"Citations without references: {missing}")
    for old_number in sorted(entries):
        if old_number not in first_order:
            first_order.append(old_number)
            first_location[old_number] = "Not cited before reference list"

    mapping = {old: new for new, old in enumerate(first_order, start=1)}

    def replace_citation(match):
        mapped = [mapping[value] for value in expand_citation(match.group(1))]
        return "[" + compress_citation(mapped) + "]"

    new_before = citation_pattern.sub(replace_citation, before_refs)
    new_reference_lines = [f"{mapping[old]}. {entries[old]}" for old in first_order]
    new_md = (
        new_before
        + "## References\n\n"
        + "\n\n".join(new_reference_lines)
        + "\n\n## Figures and legends"
        + after_refs
    )

    REF_AUDIT.parent.mkdir(parents=True, exist_ok=True)
    with REF_AUDIT.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter="\t", lineterminator="\n")
        writer.writerow([
            "old_reference_number", "new_reference_number", "first_appearance_location",
            "reference_title_short", "renumbered", "notes"
        ])
        for old in sorted(entries):
            writer.writerow([
                old, mapping[old], first_location[old], short_title(entries[old]),
                "yes", "Metadata, title, journal and DOI preserved verbatim."
            ])
    return new_md


def polish_markdown(md):
    old_background = (
        "This design links genetic priority to cellular and tissue context while preserving "
        "the distinct strength and limitations of each evidence layer (Figure 1)."
    )
    new_background = (
        "This design links genetic priority to cellular and tissue context while preserving "
        "the distinct strength and limitations of each evidence layer."
    )
    old_magma = (
        "these results define an EUR-LD-reference-based genetic-priority layer, not ancestry-"
        "generalizable fine mapping or causal-gene evidence (Figure 1)."
    )
    new_magma = (
        "these results define an EUR-LD-reference-based genetic-priority layer, not ancestry-"
        "generalizable fine mapping or causal-gene evidence (Supplementary Figure S1; "
        "Supplementary Table 2)."
    )
    if old_background not in md or old_magma not in md:
        raise ValueError("Expected Figure 1 callout text was not found in v3.1")
    md = md.replace(old_background, new_background).replace(old_magma, new_magma)

    with FIG1_AUDIT.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter="\t", lineterminator="\n")
        writer.writerow(["location", "old_text_or_issue", "new_text", "resolved", "notes"])
        writer.writerow([
            "Background final paragraph",
            "Integrative-design sentence ended with (Figure 1), although Figure 1 is GWAS QC only.",
            new_background,
            "yes",
            "Figure 1 callout removed; scientific wording otherwise unchanged."
        ])
        writer.writerow([
            "Results MAGMA paragraph",
            "EUR-LD genetic-priority statement cited Figure 1 as if MAGMA were displayed.",
            new_magma,
            "yes",
            "MAGMA display routing moved to Supplementary Figure S1 and Supplementary Table 2."
        ])
    return renumber_references(md)


def set_run_font(run, size=10, bold=False, italic=False, superscript=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript


def add_runs(paragraph, text, size=10):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\^[^\^]+\^)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        mono = part.startswith("`") and part.endswith("`")
        superscript = part.startswith("^") and part.endswith("^")
        if bold:
            clean = part[2:-2]
        elif italic or mono or superscript:
            clean = part[1:-1]
        else:
            clean = part
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, bold=bold, italic=italic, superscript=superscript)
        if mono:
            run.font.name = "Courier New"


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def update_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(header.add_run("BMC Genomics manuscript | Final polished draft"), 8)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(footer.add_run("Manuscript v3.2 final polished | Page "), 8)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def add_figure(doc, path):
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(6.45))


def build_docx(md):
    doc = Document(BASE_DOCX)
    clear_body(doc)
    update_header_footer(doc)
    figure_map = {
        "### Figure 1.": ROOT / "work/phase6_step2_docx/phase1_manhattan_figure1.png",
        "### Figure 2.": ROOT / "results/figures/phase2_step5C_Figure2_snRNA_context_final_draft_600dpi.png",
        "### Figure 3.": ROOT / "results/figures/phase6_step4_Figure3_candidate_reporting_model_bulk_reviewed_600dpi.png",
        "### Figure 4.": ROOT / "results/figures/phase5_step4_Figure4_bulk_disease_context_draft_600dpi.png",
    }
    in_references = False
    pending_supplementary_legend = False
    for block in re.split(r"\n\s*\n", md.strip()):
        line = block.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.add_run(line[2:])
            set_run_font(run, 18, bold=True)
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif line.startswith("## "):
            heading = line[3:]
            if heading == "Figures and legends":
                in_references = False
                continue
            in_references = heading == "References"
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            add_runs(paragraph, heading, 12)
        elif line.startswith("### "):
            key = next((key for key in figure_map if line.startswith(key)), None)
            if key:
                doc.add_page_break()
                add_figure(doc, figure_map[key])
            elif line.startswith("### Supplementary Figure S1."):
                doc.add_page_break()
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            add_runs(paragraph, line[4:], 11)
            pending_supplementary_legend = line.startswith("### Supplementary Figure")
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.08
            if in_references:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.first_line_indent = Inches(-0.18)
                paragraph.paragraph_format.keep_together = True
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if pending_supplementary_legend:
                paragraph.paragraph_format.keep_together = True
                pending_supplementary_legend = False
            add_runs(paragraph, line.replace("\n", " "), 10)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    doc.save(OUT_DOCX)


def main():
    md = polish_markdown(INPUT_MD.read_text(encoding="utf-8"))
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(md)


if __name__ == "__main__":
    main()
