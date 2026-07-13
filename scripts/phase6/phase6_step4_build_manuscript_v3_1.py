#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
BASE_MD = ROOT / "docs/revision/stage7K_github_url_patch/manuscript_v2.5_BMC_Genomics_GitHub_URL_inserted.md"
BASE_V30_MD = ROOT / "manuscript/manuscript_v3.0_context_mapping_revision.md"
BASE_DOCX = ROOT / "manuscript/manuscript_v3.0_context_mapping_revision.docx"
OUTDIR = ROOT / "manuscript"
WORK = ROOT / "work/phase6_step2_docx"

TITLE = "Post-GWAS renal papillary context mapping of kidney stone disease genetic risk"

ABSTRACT = """**Background:** Kidney stone disease (KSD) has a substantial polygenic component, but the renal cell and tissue contexts in which associated genes are represented remain incompletely resolved. We integrated trans-ancestry KSD genome-wide association summary statistics with renal papillary single-nucleus, spatial and paired bulk transcriptomic resources. MAGMA supplied an EUR-linkage-disequilibrium-reference-based genetic-priority layer, while GTEx Kidney_Cortex S-PrediXcan was treated as proxy annotation rather than papilla-specific regulatory evidence.

**Results:** Quality control retained 4,915,033 of 5,960,489 GWAS rows, including 3,209 genome-wide significant variants, and identified 60 lead variants in 57 reconstructed loci. MAGMA tested 17,316 genes and identified 94 Bonferroni-significant, 369 false-discovery-rate (FDR)-significant and 187 suggestive genes. In GSE231569, 43,878 nuclei from four papillary donors, including 540 Loop of Henle/thick ascending limb (Loop/TAL) nuclei, supported a donor-level Loop/TAL-associated context with partial matched-random support and persistence after known-driver removal. Spatial label transfer across ten GSE206306 sections was usable only for broad-compartment projection; Loop/TAL predictions were zero or sparse, precluding claim-grade Loop/TAL or lesion enrichment. Kidney_Cortex S-PrediXcan tested 5,989 genes and yielded 51 FDR-supported models, of which 42 were one-SNP and nine were multi-SNP models. A repaired 235-gene reporting set comprised R1-R6 groups of 68, 1, 2, 141, 16 and 7 genes, respectively. In GSE73680, 55 samples from 29 patients, including 26 paired patients, showed concordant positive module coefficients that did not remain FDR-significant and were attenuated by tissue-state sensitivity analyses.

**Conclusions:** Genetically prioritized KSD genes map to a donor-level Loop/TAL-associated renal papillary context and are embedded in a broader injury/remodeling-associated bulk disease background. These results provide post-GWAS functional interpretation and context mapping, not causal gene validation, causal cell-type assignment, plaque-specific localization or papilla-specific regulatory inference."""

RESULTS = """### Auditable GWAS and MAGMA analyses define an EUR-LD-reference-based genetic-priority layer

The locked GWAS input contained 5,960,489 rows. Quality control retained 4,915,033 rows, including 3,209 variants at genome-wide significance. Distance pruning identified 60 lead SNPs, and merging overlapping windows produced 57 reconstructed loci (Supplementary Table 1). The source publication reported 59 loci,[4] but the publication-level lead-SNP or locus table required to reconcile the two counts was not available. We therefore report the 57-locus result as a reproducible reconstruction rather than as a replacement for the published locus definition.

MAGMA v1.10 tested 17,316 genes using GRCh37 gene locations and the 1000 Genomes European reference panel. Ninety-four genes met Bonferroni significance, 369 met FDR significance and 187 met the prespecified suggestive threshold of P < 1 x 10^-4 (Supplementary Table 2). Locked top-50, top-100, Bonferroni, FDR and suggestive modules were propagated without post hoc membership changes. Because the source GWAS was trans-ancestry whereas the linkage-disequilibrium reference was European, these results define an EUR-LD-reference-based genetic-priority layer, not ancestry-generalizable fine mapping or causal-gene evidence (Figure 1).

### Donor-level single-nucleus analyses map prioritized modules to a Loop/TAL-associated context

The audited GSE231569 papillary single-nucleus dataset contained 43,878 nuclei from four donors. Harmonized broad-compartment annotation identified 540 nuclei in the Loop/TAL compartment. Donor-by-compartment summaries placed the locked MAGMA modules consistently toward the Loop/TAL end of the expression distribution, with donors serving as the biological support units rather than individual nuclei (Figure 2; Supplementary Table 3).

The direction of the Loop/TAL-versus-other-compartment contrast remained stable in leave-one-donor-out analyses. Expression- and detection-matched random gene sets provided partial, rather than uniform, separation from matched expectation, and rank statistics showed ceiling effects. Panel-level removal of prespecified Loop/TAL and renal-transport drivers did not erase the overall direction. Taken together, these analyses support a moderate donor-level Loop/TAL-associated context. They do not establish Loop/TAL as a unique or causal cell type for KSD genetic effects.

### Spatial data provide descriptive broad-compartment tissue-context projection

Ten complete GSE206306 spatial sections were retained for descriptive projection. Single-nucleus-to-spatial label transfer was usable for broad-compartment orientation, although two sections showed lower confidence. Loop/TAL predictions were zero or sparse across sections, so the spatial data could not support Loop/TAL enrichment, Loop/TAL co-distribution or cross-section replication claims. No plaque, mineral, calcification, fibrosis or lesion region-of-interest annotation was available.

Locked MAGMA module scores were overlaid descriptively on the tissue sections and interpreted only against broad transferred compartments (Supplementary Figure S2). These overlays show where module signal occurs within the assayed papillary tissue, but they do not localize genetic risk to plaque or infer disease-control differences, papilla-specific regulation or causal anatomy.

### A repaired candidate model separates MAGMA priority from Kidney_Cortex proxy annotation

S-PrediXcan with GTEx v8 Kidney_Cortex models tested 5,989 genes and identified 51 FDR-supported models. Forty-two were one-SNP models and nine were multi-SNP models, making the evidence predominantly weak and variant-limited. Kidney_Cortex also does not match the renal papilla. TWAS was therefore retained as proxy annotation and was not used to upgrade genes to papilla-specific regulatory or causal status (Supplementary Figure S3; Supplementary Table 4).

The repaired candidate universe contained 235 genes assigned deterministically to six mutually exclusive reporting groups: R1, 68; R2, 1; R3, 2; R4, 141; R5, 16; and R6, 7 (Figure 3; Supplementary Table 5). These are reporting groups, not causal tiers. R1-R6 organize combinations of MAGMA priority, proxy TWAS support and contextual review; curated exemplars remain interpretive annotations and do not alter assignment. Complete resources for claim-grade SMR/HEIDI or colocalization were unavailable, and no such result was used.

### Paired bulk analyses place prioritized modules in an injury/remodeling-associated disease context

The reconstructed GSE73680 dataset comprised 55 papillary samples from 29 patients, including 26 patients with paired control/adjacent and plaque/stone-associated tissue. Module scores were calculated as the arithmetic mean of gene-wise z scores across detected members, without an additional log transformation. In patient fixed-effect models, all five locked MAGMA modules had positive group coefficients (0.169-0.174), but none remained significant after FDR correction (FDR 0.071-0.082) (Figure 4; Supplementary Table 6).

The paired module changes tracked broader tissue-state signatures. Injury/remodeling showed the strongest nominal association (beta = 0.375, P = 0.0057, FDR = 0.057); epithelial and mineralization/remodeling proxies were weaker (beta = 0.348, P = 0.0318, FDR = 0.106; and beta = 0.292, P = 0.0210, FDR = 0.105, respectively). Adjustment and paired-delta sensitivity analyses attenuated the module effects. Because these signatures are bulk expression scores and not estimated cell fractions, the results indicate embedding within an injury/remodeling-associated disease context, not cell-composition-independent regulation, deconvolution or causal validation.

### Integrated evidence supports context mapping while preserving inferential boundaries

Across layers, MAGMA provided the genetic-priority anchor; donor-level snRNA data supplied the clearest renal cell-context signal; spatial data contributed supplementary broad-compartment anatomy; Kidney_Cortex TWAS supplied limited proxy annotation; and paired bulk data linked the modules to a broader disease-associated tissue state. The layers are complementary but not interchangeable, and agreement between them does not constitute causal triangulation.

The integrated result is therefore a post-GWAS renal papillary context map. It prioritizes a Loop/TAL-associated single-nucleus context and an attenuated injury/remodeling-associated bulk background while explicitly leaving causal genes, causal cell populations, plaque-specific localization and papilla-specific regulatory mechanisms unresolved."""

DISCUSSION = """The principal contribution of this study is an evidence-stratified map connecting KSD gene-based association signals to renal papillary transcriptomic contexts. Rather than treating overlap across datasets as validation, the analysis assigns each layer a defined inferential role. MAGMA anchors genetic priority, donor-level snRNA identifies a Loop/TAL-associated context, spatial data provide descriptive anatomical orientation, Kidney_Cortex TWAS supplies proxy annotation and paired bulk data capture the disease-associated tissue background.

The Loop/TAL-associated single-nucleus pattern is biologically compatible with renal transport and ion-handling pathways implicated in stone susceptibility. Established genes including *UMOD*, *CASR*, *CLDN14* and *CLDN10* provide useful biological anchors.[18-23] Their familiarity, however, cannot identify the gene mediating a GWAS locus. The persistence of the module-level pattern after driver removal makes the result less dependent on a small set of canonical markers, while the partial matched-random benchmark and four-donor design appropriately limit its strength to moderate contextual support.

The bulk results provide a complementary tissue-state view. Concordant positive coefficients across five modules were not FDR-significant and were weakened by injury/remodeling and epithelial-state sensitivity analyses. This attenuation is informative because it places genetically prioritized programs within a broader papillary response rather than isolating a disease-specific module. It does not distinguish transcriptional regulation within cells from shifts in cellular composition.

The spatial and TWAS layers chiefly define boundaries. Ten spatial sections permitted broad-compartment projection, but sparse or absent Loop/TAL transfer and missing lesion annotations prevented plaque- or Loop/TAL-specific inference. Likewise, the predominance of one-SNP Kidney_Cortex models and the cortex-to-papilla mismatch prevented regulatory upgrading of individual genes. Neither missing claim-grade SMR/colocalization nor weak proxy support should be interpreted as evidence against biological relevance; they identify experiments and resources still needed.

Future work should combine ancestry-matched genetic analyses, locus-level fine mapping and papilla-relevant eQTL resources with lesion-resolved spatial profiling. Donor-aware perturbation in relevant tubular systems could then test whether selected genes and modules exert direct effects or mark downstream tissue responses. These steps are required to move from contextual prioritization to mechanism.

## Limitations

The use of a 1000 Genomes European LD reference for a trans-ancestry GWAS limits ancestry-generalizable interpretation. The 57 reconstructed loci could not be reconciled with the 59 published loci without the source lead-SNP or locus table. The single-nucleus analysis contained four donors and only 540 Loop/TAL nuclei, and matched-random support was partial. Spatial projection was descriptive, included two lower-confidence sections and lacked usable Loop/TAL prediction and lesion annotations. Kidney_Cortex TWAS was tissue-mismatched and dominated by one-SNP models, while complete inputs for claim-grade SMR/HEIDI and colocalization were unavailable. The bulk dataset was reconstructed from microarray data; its tissue-state scores were not cell fractions, and module associations did not survive FDR correction. All analyses were observational, and no perturbational experiment tested a prioritized gene, module or cell context."""

METHODS = """### Study design and evidence hierarchy

The study integrated the 2025 trans-ancestry KSD GWAS,[4] GSE231569 papillary snRNA data,[15] ten GSE206306 spatial sections,[15] GSE73680 paired papillary bulk expression,[11] and GTEx v8 Kidney_Cortex prediction models.[35] Analyses were assigned prospectively by inferential role: GWAS/MAGMA for genetic priority, snRNA for donor-level cell context, spatial data for supplementary broad-compartment projection, Kidney_Cortex TWAS for proxy annotation and bulk data for paired disease-context sensitivity.

### GWAS quality control and locus reconstruction

GRCh37-compatible summary statistics were restricted to valid biallelic A/C/G/T SNP records with valid coordinates and P values; strand-ambiguous variants and duplicate or malformed records were removed, with low-frequency filtering applied where allele frequency was available. Genome-wide significant variants were greedily pruned within +/-1 Mb and overlapping lead windows were merged. The procedure retained 4,915,033 of 5,960,489 rows, identified 3,209 genome-wide significant variants and 60 lead SNPs, and reconstructed 57 loci. The reconstruction was not used to explain the difference from the 59 loci reported by the source study.

### MAGMA analysis and canonical module freeze

MAGMA v1.10 used NCBI build 37 gene locations and the 1000 Genomes European LD reference.[24] Of 17,316 tested genes, significance was defined using Bonferroni correction, Benjamini-Hochberg FDR,[34] and a suggestive threshold of P < 1 x 10^-4. Top-50, top-100, Bonferroni, FDR and suggestive lists were frozen before downstream scoring. Executable, version, gene-location, LD-prefix, SNP/P input, gene output and log provenance were audited. MAGMA results were interpreted as EUR-LD-reference-based genetic priority.

### Single-nucleus processing, scoring and sensitivity analyses

Four papillary donors from GSE231569 were retained, yielding 43,878 audited nuclei after quality control and broad-compartment harmonization. The Loop/TAL compartment contained 540 nuclei. For each frozen module, nucleus-level scores were calculated as arithmetic means of log-normalized expression among detected genes and summarized by donor and compartment. Donors were biological support units. Leave-one-donor-out summaries, expression- and detection-matched random gene sets, and panel-level removal of prespecified Loop/TAL/transport drivers evaluated donor dependence, matched expectation and driver dependence.

### Spatial broad-compartment projection and descriptive module overlays

Ten complete GSE206306 sections with expression, coordinate and image assets were processed as spatial objects. The GSE231569 reference supported label transfer only at the broad-compartment level. Transfer confidence and prediction sparsity were audited by section. Because Loop/TAL calls were zero or sparse and no lesion-region annotations were available, module scores were displayed as descriptive overlays without formal Loop/TAL, plaque, disease-control or cross-section enrichment tests. Sections, not spots, defined the highest available biological unit.

### Kidney_Cortex proxy TWAS and candidate reporting groups

S-PrediXcan used GTEx v8 Kidney_Cortex MASHR models.[35,36] Benjamini-Hochberg correction was applied across 5,989 tested genes, and 51 FDR-supported models were classified by the number of contributing variants as one-SNP (42) or multi-SNP (9). The tissue was treated as a cortex proxy, and TWAS did not upgrade causal or papilla-specific regulatory status. A deterministic union and deduplication procedure produced 235 candidate genes and mutually exclusive R1-R6 reporting groups (68, 1, 2, 141, 16 and 7 genes). The groups are reporting categories, not causal tiers. Feasibility of SMR/HEIDI and colocalization was audited, but incomplete inputs precluded claim-grade analyses.

### Paired bulk reconstruction, module scoring and tissue-state sensitivity

GSE73680 feature-level expression and curated metadata yielded 55 samples from 29 patients, including 26 complete pairs.[11,32] Module scores were arithmetic means of gene-wise z scores across the 55 samples for detected module genes. Paired deltas were defined as plaque/stone-associated minus control/adjacent, and patient fixed-effect models used `module_score ~ group_binary + factor(patient_id)`. Benjamini-Hochberg correction was applied across the five canonical modules.

Injury/remodeling, epithelial, extracellular-matrix/fibrosis, immune/inflammatory and mineralization/remodeling marker scores were used as tissue-state expression proxies. Paired-delta and patient-aware covariate models tested whether module effects persisted after adjustment. These scores were not interpreted as cell fractions, and the analyses were not described as deconvolution or mediation.

### Statistical reporting and reproducibility

Tests were two-sided where applicable, with correction families defined within each analysis layer. Donors, patients and sections, rather than nuclei, samples or spots, were treated as biological support units for cross-unit interpretation. Machine-readable source data, manifests, logs and scripts were retained for reported tables and figures. No exploratory SMR, colocalization, plaque localization or spatial Loop/TAL analysis was promoted to claim-grade evidence."""

FIG_LEGENDS = """## Figures and legends

### Figure 1. Auditable GWAS and EUR-LD-reference-based MAGMA genetic priority

Manhattan representation of the locked trans-ancestry KSD GWAS after quality control. The analysis retained 4,915,033 variants, including 3,209 genome-wide significant variants, and identified 60 lead SNPs in 57 reconstructed loci. MAGMA tested 17,316 genes using the 1000 Genomes European LD reference; 94, 369 and 187 genes met Bonferroni, FDR and suggestive criteria, respectively. The figure establishes a genetic-priority layer and does not imply ancestry-generalizable fine mapping or causal genes.

### Figure 2. Donor-level snRNA evidence for a Loop/TAL-associated context

Donor-aware module summaries in 43,878 nuclei from four GSE231569 papillary donors, including 540 Loop/TAL nuclei. Panels show broad-compartment context, donor-level module contrasts, matched-random benchmarking, leave-one-donor-out sensitivity and known-driver removal. Donors are the biological support units. The figure supports a moderate Loop/TAL-associated context, not a causal cell-type assignment.

### Figure 3. Repaired candidate reporting model

Evidence organization for 235 genes across six mutually exclusive reporting groups: R1 = 68, R2 = 1, R3 = 2, R4 = 141, R5 = 16 and R6 = 7. MAGMA priority is separated from GTEx Kidney_Cortex proxy TWAS support, including one-SNP and multi-SNP model classes. R1-R6 are reporting groups rather than causal tiers, and curated exemplars do not alter group assignment.

### Figure 4. Paired bulk disease-context analysis

GSE73680 analysis of 55 samples from 29 patients, including 26 complete pairs. Five canonical MAGMA module coefficients were positive but did not remain FDR-significant. Tissue-state and paired-delta sensitivity analyses show attenuation and association with a broader injury/remodeling background. Marker scores are expression proxies, not cell fractions.

### Supplementary Figure S1. GWAS and MAGMA diagnostic plots

Locked GWAS Manhattan and quantile-quantile diagnostics, P-value distribution and MAGMA gene-rank summaries. These panels document input behavior and gene-priority thresholds; they do not provide locus fine mapping or causal-gene inference.

### Supplementary Figure S2. Descriptive spatial broad-compartment projection

Broad-compartment label-transfer quality control and descriptive locked-module overlays across ten GSE206306 sections. Two sections showed lower confidence, Loop/TAL predictions were zero or sparse and no lesion-region annotation was available. The panels do not support plaque localization, Loop/TAL enrichment, disease-control comparison or causal spatial inference.

### Supplementary Figure S3. Full Kidney_Cortex TWAS proxy and candidate evidence display

Complete S-PrediXcan and candidate evidence summaries for 5,989 tested models, including 51 FDR-supported results (42 one-SNP and nine multi-SNP models), together with the 235-gene reporting universe. The figure presents cortex-proxy annotation and reporting categories, not papilla-specific regulation or causal tiers."""


def section(text, start, end=None):
    i = text.index(start)
    j = text.index(end, i) if end else len(text)
    return text[i:j].strip()


def build_markdown():
    old = BASE_MD.read_text(encoding="utf-8")
    author = section(old, "**Article type:**", "## Abstract")
    background = section(old, "## Background", "## Results")
    old_close = ("Here, we developed an evidence-stratified post-GWAS framework to determine whether MAGMA-prioritized KSD genes map to a coherent renal papillary context. We separated MAGMA genetic priority from Kidney_Cortex TWAS proxy support, evaluated donor-level single-nucleus patterns in GSE231569, tested paired bulk module responses in GSE73680 with composition- and injury/remodeling-aware sensitivity analyses, and assessed complexity-adjusted spatial co-distribution in GSE206306. This design links genetic priority to cellular and tissue context while preserving the distinct strength and limitations of each evidence layer (Figure 1).")
    new_close = ("Here, we developed an evidence-stratified post-GWAS framework to determine whether MAGMA-prioritized KSD genes map to coherent renal papillary contexts. We separated EUR-LD-reference-based MAGMA genetic priority from Kidney_Cortex TWAS proxy annotation, evaluated donor-level single-nucleus patterns in GSE231569, tested paired bulk module responses in GSE73680 with tissue-state sensitivity analyses, and used GSE206306 only for descriptive broad-compartment spatial projection. This design links genetic priority to cellular and tissue context while preserving the distinct strength and limitations of each evidence layer (Figure 1).")
    background = background.replace(old_close, new_close)
    declarations = section(old, "## Declarations", "## References")
    declarations = declarations.replace(
        "and will be archived at Zenodo before submission [TO FILL: repository Zenodo DOI].",
        "and will be archived at Zenodo upon release (Zenodo DOI to be added upon release).",
    )
    references = section(old, "## References", "## Figure legends")
    return f"""# {TITLE}\n\n{author}\n\n## Abstract\n\n{ABSTRACT}\n\n## Keywords\n\nkidney stone disease; genome-wide association study; MAGMA; renal papilla; Loop of Henle; thick ascending limb; single-nucleus RNA sequencing; spatial transcriptomics; TWAS\n\n{background}\n\n## Results\n\n{RESULTS}\n\n## Discussion\n\n{DISCUSSION}\n\n## Methods\n\n{METHODS}\n\n## Conclusions\n\nGenetically prioritized KSD genes map to a donor-level Loop/TAL-associated renal papillary context and an attenuated injury/remodeling-associated bulk disease background. This framework provides post-GWAS functional interpretation while preserving the distinction between contextual support and causal gene, cell-type, lesion-localization or papilla-specific regulatory evidence.\n\n## List of abbreviations\n\nBH: Benjamini-Hochberg; ECM: extracellular matrix; eQTL: expression quantitative trait locus; FDR: false discovery rate; GEO: Gene Expression Omnibus; GWAS: genome-wide association study; KSD: kidney stone disease; LD: linkage disequilibrium; MAGMA: Multi-marker Analysis of GenoMic Annotation; QC: quality control; ROI: region of interest; SMR: summary-data Mendelian randomization; SNP: single-nucleotide polymorphism; snRNA-seq: single-nucleus RNA sequencing; TAL: thick ascending limb; TWAS: transcriptome-wide association study; UMAP: uniform manifold approximation and projection.\n\n{declarations}\n\n{references}\n\n{FIG_LEGENDS}\n"""


def targeted_revision(md):
    def rr(old, new):
        nonlocal md
        if old not in md:
            raise ValueError(f"v3.1 source text not found: {old[:100]}")
        md = md.replace(old, new)

    rr(
        "supported a donor-level Loop/TAL-associated context with partial matched-random support and persistence after known-driver removal. Spatial label transfer across ten GSE206306 sections was usable only for broad-compartment projection",
        "supported a donor-level Loop/TAL-associated context in expression- and detection-matched random benchmarks, with attenuation but retention of the broader pattern after known-driver removal. Spatial label transfer across ten complete sections, comprising four GSE206306 sections and six GSE231630 sections, was usable only for broad-compartment projection",
    )
    rr("were attenuated by tissue-state sensitivity analyses.", "were attenuated in sensitivity analyses using a KIM1/LCN2-like injury proxy and related tissue-state scores.")
    rr("are embedded in a broader injury/remodeling-associated bulk disease background.", "are embedded in a broader KIM1/LCN2-like injury-associated bulk disease context.")
    rr("and used GSE206306 only for descriptive broad-compartment spatial projection.", "and used four GSE206306 and six GSE231630 sections only for descriptive broad-compartment spatial projection.")
    rr(
        "Expression- and detection-matched random gene sets provided partial, rather than uniform, separation from matched expectation, and rank statistics showed ceiling effects. Panel-level removal of prespecified Loop/TAL and renal-transport drivers did not erase the overall direction.",
        "Expression- and detection-matched random benchmarks supported the primary donor-level Loop/TAL-minus-other contrast across canonical modules, whereas rank statistics showed ceiling effects. Panel-level removal of prespecified Loop/TAL and renal-transport drivers attenuated but did not abolish the broader module-level pattern; retention after removal does not prove independence from Loop/TAL biology.",
    )
    rr("Ten complete GSE206306 spatial sections were retained for descriptive projection.", "Ten complete spatial sections were retained for descriptive projection, comprising four GSE206306 sections and six GSE231630 sections.")
    rr(
        "curated exemplars remain interpretive annotations and do not alter assignment.",
        "curated exemplars remain interpretive annotations and do not alter assignment. Bulk status indicates paired disease-context and tissue-state review and is not used to upgrade candidate reporting groups.",
    )
    rr("### Paired bulk analyses place prioritized modules in an injury/remodeling-associated disease context", "### Paired bulk analyses place prioritized modules in a KIM1/LCN2-like injury-associated disease context")
    rr(
        "The reconstructed GSE73680 dataset comprised 55 papillary samples from 29 patients, including 26 patients with paired control/adjacent and plaque/stone-associated tissue.",
        "The reconstructed GSE73680 dataset comprised 55 papillary samples from 29 patients, including 26 patients with paired control/adjacent and plaque/stone-associated tissue. Group labels were filename-derived and manually reviewable.",
    )
    rr("Injury/remodeling showed the strongest nominal association", "The KIM1/LCN2-like injury proxy showed the strongest nominal positive shift")
    rr("embedding within an injury/remodeling-associated disease context", "embedding within a KIM1/LCN2-like injury-associated tissue-state context")
    rr("an attenuated injury/remodeling-associated bulk background", "an attenuated KIM1/LCN2-like injury-associated bulk background")
    rr("an attenuated injury/remodeling-associated bulk disease background", "an attenuated KIM1/LCN2-like injury-associated bulk disease context")
    rr(
        "The persistence of the module-level pattern after driver removal makes the result less dependent on a small set of canonical markers, while the partial matched-random benchmark and four-donor design appropriately limit its strength to moderate contextual support.",
        "Expression- and detection-matched random benchmarks supported the donor-level Loop/TAL-minus-other contrast across canonical modules. Known-driver removal attenuated but did not abolish the broader module-level pattern; this sensitivity does not establish independence from Loop/TAL biology. The four-donor design and one donor with only four Loop/TAL nuclei appropriately limit the result to moderate contextual support.",
    )
    rr("were weakened by injury/remodeling and epithelial-state sensitivity analyses.", "were attenuated by KIM1/LCN2-like injury-proxy and epithelial-state sensitivity analyses.")
    rr("Ten spatial sections permitted broad-compartment projection", "Ten spatial sections from GSE206306 and GSE231630 permitted broad-compartment projection")
    rr("and matched-random support was partial.", "and one donor contributed only four Loop/TAL nuclei. Matched-random benchmarks supported the canonical contrasts, whereas known-driver removal attenuated but did not abolish the broader pattern.")
    rr(
        "The bulk dataset was reconstructed from microarray data; its tissue-state scores were not cell fractions",
        "The bulk dataset was reconstructed from microarray data, and its group labels were filename-derived and manually reviewable; its tissue-state scores were not cell fractions",
    )
    rr("GSE231569 papillary snRNA data,[15] ten GSE206306 spatial sections,[15] GSE73680", "GSE231569 papillary snRNA data,[15] ten spatial sections comprising four GSE206306 and six GSE231630 sections,[15] GSE73680")
    rr(
        "Four papillary donors from GSE231569 were retained, yielding 43,878 audited nuclei after quality control and broad-compartment harmonization.",
        "Four papillary donors from GSE231569 were retained, yielding 43,878 audited nuclei after quality control and broad-compartment harmonization. Seurat was used for single-nucleus object handling, normalization and broad-compartment workflows.[28]",
    )
    rr("Leave-one-donor-out summaries, expression- and detection-matched random gene sets, and panel-level removal", "Leave-one-donor-out summaries, 1,000 expression- and detection-matched random gene sets per module, low-cell-count sensitivity, and panel-level removal")
    rr("Ten complete GSE206306 sections with expression, coordinate and image assets were processed as spatial objects.", "Ten complete sections with expression, coordinate and image assets were processed as Seurat spatial objects, comprising four GSE206306 and six GSE231630 sections.[28]")
    rr("Because Loop/TAL calls were zero or sparse", "Because transferred Loop/TAL prediction scores were zero or nonzero-sparse")
    rr("Feasibility of SMR/HEIDI and colocalization was audited", "Feasibility of SMR/HEIDI[30] and colocalization[31,38] was audited")
    rr("GSE73680 feature-level expression and curated metadata yielded 55 samples", "GSE73680 feature-level expression and filename-derived, manually reviewable group labels yielded 55 samples")
    rr("Injury/remodeling, epithelial, extracellular-matrix/fibrosis", "KIM1/LCN2-like injury, epithelial, extracellular-matrix/fibrosis")
    rr("under accessions GSE231569, GSE73680 and GSE206306.", "under accessions GSE231569, GSE73680, GSE206306 and GSE231630.")

    rr("### Figure 1. Auditable GWAS and EUR-LD-reference-based MAGMA genetic priority", "### Figure 1. GWAS quality-control Manhattan plot")
    rr(
        "Manhattan representation of the locked trans-ancestry KSD GWAS after quality control. The analysis retained 4,915,033 variants, including 3,209 genome-wide significant variants, and identified 60 lead SNPs in 57 reconstructed loci. MAGMA tested 17,316 genes using the 1000 Genomes European LD reference; 94, 369 and 187 genes met Bonferroni, FDR and suggestive criteria, respectively. The figure establishes a genetic-priority layer and does not imply ancestry-generalizable fine mapping or causal genes.",
        "Downsampled Manhattan representation of the locked trans-ancestry KSD GWAS after quality control, with all sampled genome-wide significant variants retained for visualization. The plot documents GWAS input behavior only. MAGMA gene counts and rank summaries are reported in the Results, Supplementary Figure S1 and Supplementary Tables; they are not displayed in this panel. The figure does not imply ancestry-generalizable fine mapping or causal genes.",
    )
    rr("The figure supports a moderate Loop/TAL-associated context, not a causal cell-type assignment.", "Expression- and detection-matched random benchmarks support the donor-level contrast, whereas known-driver removal attenuates but does not abolish the broader pattern. The figure supports a moderate Loop/TAL-associated context, not causal cell-type assignment or independence from Loop/TAL biology.")
    rr("curated exemplars do not alter group assignment.", "curated exemplars do not alter group assignment. Bulk status indicates paired disease-context and tissue-state review and is not used to upgrade candidate reporting groups.")
    rr("association with a broader injury/remodeling background. Marker scores are expression proxies", "association with a broader KIM1/LCN2-like injury-proxy background. Group labels were filename-derived and manually reviewable; marker scores are expression proxies")
    rr("Locked GWAS Manhattan and quantile-quantile diagnostics, P-value distribution and MAGMA gene-rank summaries.", "Quantile-quantile diagnostics, P-value distribution and MAGMA gene-rank summaries complement the main-text GWAS QC Manhattan plot.")
    rr("across ten GSE206306 sections.", "across ten sections comprising four GSE206306 and six GSE231630 sections.")

    # limma was not used in the locked bulk workflow. Remove its unused reference,
    # shift later reference numbers, and update numeric citations deterministically.
    md = re.sub(r"\n27\. Ritchie ME,.*?https://doi\.org/10\.1093/nar/gkv007\n", "\n", md)
    refs_start = md.index("## References")
    head, refs = md[:refs_start], md[refs_start:]
    refs = re.sub(
        r"(?m)^(\d+)\. ",
        lambda m: f"{int(m.group(1)) - 1 if int(m.group(1)) >= 28 else int(m.group(1))}. ",
        refs,
    )

    def shift_citation(match):
        out = []
        for token in match.group(1).split(","):
            token = token.strip()
            if "-" in token:
                a, b = map(int, token.split("-"))
                out.append(f"{a-1 if a >= 28 else a}-{b-1 if b >= 28 else b}")
            else:
                n = int(token)
                out.append(str(n - 1 if n >= 28 else n))
        return "[" + ",".join(out) + "]"

    head = re.sub(r"\[([0-9, -]+)\]", shift_citation, head)
    return head + refs


def set_cell_text_font(run, size=10, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_runs(paragraph, text, size=10):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        mono = part.startswith("`") and part.endswith("`")
        clean = part[2:-2] if bold else part[1:-1] if (italic or mono) else part
        run = paragraph.add_run(clean)
        set_cell_text_font(run, size=size, bold=bold, italic=italic)
        if mono:
            run.font.name = "Courier New"


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def render_figure1():
    out = WORK / "phase1_manhattan_figure1.png"
    if not out.exists():
        raise FileNotFoundError(
            "Render results/figures/phase1_step4_gwas_manhattan_plot.pdf "
            "to work/phase6_step2_docx/phase1_manhattan_figure1.png before building."
        )
    return out


def add_figure(doc, path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.45))


def update_header_footer(doc):
    sec = doc.sections[0]
    for p in sec.header.paragraphs:
        p.text = "BMC Genomics manuscript | Context-mapping revision"
        for r in p.runs:
            set_cell_text_font(r, 8)
    fp = sec.footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Manuscript v3.1 targeted revision | Page ")
    set_cell_text_font(r, 8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def build_docx(md):
    doc = Document(BASE_DOCX)
    clear_body(doc)
    update_header_footer(doc)
    fig_map = {
        "### Figure 1.": render_figure1(),
        "### Figure 2.": ROOT / "results/figures/phase2_step5C_Figure2_snRNA_context_final_draft_600dpi.png",
        "### Figure 3.": ROOT / "results/figures/phase6_step4_Figure3_candidate_reporting_model_bulk_reviewed_600dpi.png",
        "### Figure 4.": ROOT / "results/figures/phase5_step4_Figure4_bulk_disease_context_draft_600dpi.png",
    }
    first = True
    for block in re.split(r"\n\s*\n", md.strip()):
        line = block.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(line[2:])
            set_cell_text_font(r, 18, bold=True)
            r.font.color.rgb = RGBColor(31, 78, 121)
            first = False
        elif line.startswith("## "):
            if line == "## Figures and legends":
                continue
            p = doc.add_paragraph(style="Heading 1")
            add_runs(p, line[3:], 12)
        elif line.startswith("### "):
            key = next((k for k in fig_map if line.startswith(k)), None)
            if key:
                doc.add_page_break()
                add_figure(doc, fig_map[key])
            p = doc.add_paragraph(style="Heading 2")
            add_runs(p, line[4:], 11)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08
            add_runs(p, line.replace("\n", " "), 10)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
    out = OUTDIR / "manuscript_v3.1_targeted_revision.docx"
    doc.save(out)


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    WORK.mkdir(parents=True, exist_ok=True)
    md = targeted_revision(BASE_V30_MD.read_text(encoding="utf-8"))
    (OUTDIR / "manuscript_v3.1_targeted_revision.md").write_text(md, encoding="utf-8")
    build_docx(md)


if __name__ == "__main__":
    main()
